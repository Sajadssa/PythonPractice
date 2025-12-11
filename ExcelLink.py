import os
import re
from pathlib import Path

def extract_date_from_pdf(pdf_path):
    """استخراج Date از جدول در صفحه اول PDF"""
    try:
        import PyPDF2
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            first_page = pdf_reader.pages[0]
            text = first_page.extract_text()
            
            # تلاش برای یافتن Date در ساختار جدول
            lines = text.split('\n')
            for i, line in enumerate(lines):
                if 'date' in line.lower():
                    # بررسی همان خط
                    patterns = [
                        r'Date[:\s]*(\d{1,2}[-/]\w{3}[-/]\d{2,4})',
                        r'Date[:\s]*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',
                        r'Date[:\s]*(\w{3}[-/]\d{1,2}[-/]\d{2,4})',
                    ]
                    for pattern in patterns:
                        match = re.search(pattern, line, re.IGNORECASE)
                        if match:
                            return match.group(1)
                    
                    # بررسی خط بعدی
                    if i + 1 < len(lines):
                        next_line = lines[i + 1].strip()
                        date_patterns = [
                            r'^(\d{1,2}[-/]\w{3}[-/]\d{2,4})$',
                            r'^(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})$',
                            r'^(\w{3}[-/]\d{1,2}[-/]\d{2,4})$',
                        ]
                        for pattern in date_patterns:
                            match = re.match(pattern, next_line)
                            if match:
                                return match.group(1)
            
            # جستجوی کلی در متن
            general_patterns = [
                r'(\d{1,2}[-/]\w{3}[-/]\d{2,4})',
                r'(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',
            ]
            for pattern in general_patterns:
                match = re.search(pattern, text)
                if match:
                    return match.group(1)
            
            return "N/A"
            
    except ImportError:
        print("⚠️ کتابخانه PyPDF2 نصب نیست. برای نصب: pip install PyPDF2")
        return "N/A"
    except Exception as e:
        print(f"⚠️ خطا در خواندن فایل {os.path.basename(pdf_path)}: {e}")
        return "N/A"

def extract_date_from_word(doc_path):
    """استخراج Date از فایل Word - از جدول صفحه اول"""
    try:
        import docx
        doc = docx.Document(doc_path)
        
        # جستجو در جداول
        for table in doc.tables:
            for row in table.rows:
                cells_text = [cell.text.strip() for cell in row.cells]
                
                # جستجوی سلول Date و سلول بعدی آن
                for i, cell_text in enumerate(cells_text):
                    if 'date' in cell_text.lower() and i + 1 < len(cells_text):
                        date_value = cells_text[i + 1].strip()
                        if date_value and date_value.lower() != 'date':
                            return date_value
                    
                    # یا اگر Date و تاریخ در یک سلول باشند
                    if 'date' in cell_text.lower():
                        patterns = [
                            r'Date[:\s]*(\d{1,2}[-/]\w{3}[-/]\d{2,4})',
                            r'Date[:\s]*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',
                            r'Date[:\s]*(\w{3}[-/]\d{1,2}[-/]\d{2,4})',
                        ]
                        for pattern in patterns:
                            match = re.search(pattern, cell_text, re.IGNORECASE)
                            if match:
                                return match.group(1)
        
        # جستجو در پاراگراف‌ها
        for para in doc.paragraphs:
            text = para.text
            if 'date' in text.lower():
                patterns = [
                    r'Date[:\s]*(\d{1,2}[-/]\w{3}[-/]\d{2,4})',
                    r'Date[:\s]*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',
                    r'Date[:\s]*(\w{3}[-/]\d{1,2}[-/]\d{2,4})',
                ]
                for pattern in patterns:
                    match = re.search(pattern, text, re.IGNORECASE)
                    if match:
                        return match.group(1)
        
        return "N/A"
        
    except ImportError:
        print("⚠️ کتابخانه python-docx نصب نیست. برای نصب: pip install python-docx")
        return "N/A"
    except Exception as e:
        print(f"⚠️ خطا در خواندن فایل Word {os.path.basename(doc_path)}: {e}")
        return "N/A"

def extract_report_title_from_pdf(pdf_path):
    """بررسی وجود عنوان گزارش در PDF"""
    try:
        import PyPDF2
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            first_page = pdf_reader.pages[0]
            text = first_page.extract_text()
            
            if "PRODUCTION ENGINEERING MONTHLY REPORT" in text.upper():
                return "PRODUCTION ENGINEERING MONTHLY REPORT"
            elif "JCTION ENGINEERING MONTHLY REPORT" in text.upper():
                return "JCTION ENGINEERING MONTHLY REPORT"
            else:
                return "N/A"
                
    except Exception:
        return "N/A"

def extract_report_title_from_word(doc_path):
    """بررسی وجود عنوان گزارش در Word"""
    try:
        import docx
        doc = docx.Document(doc_path)
        
        full_text = ""
        for para in doc.paragraphs:
            full_text += para.text.upper() + " "
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += cell.text.upper() + " "
        
        if "PRODUCTION ENGINEERING MONTHLY REPORT" in full_text:
            return "PRODUCTION ENGINEERING MONTHLY REPORT"
        elif "JCTION ENGINEERING MONTHLY REPORT" in full_text:
            return "JCTION ENGINEERING MONTHLY REPORT"
        else:
            return "N/A"
            
    except Exception:
        return "N/A"

def extract_report_no_from_filename(filename):
    """استخراج شماره گزارش از نام فایل"""
    name_without_ext = os.path.splitext(filename)[0]
    
    # الگوی کامل
    full_pattern = r'([A-Z0-9]+-[A-Z0-9]+-[A-Z0-9]+-[A-Z0-9]+-[0-9]+-[A-Z0-9]+-[0-9]+-[A-Z0-9]+)'
    match = re.search(full_pattern, name_without_ext)
    if match:
        return match.group(1)
    
    return name_without_ext

def create_excel_report(folder_path, output_file="Weekly_Reports.xlsx"):
    """ایجاد فایل اکسل با لینک فایل‌ها و اطلاعات استخراج شده"""
    
    print("=" * 60)
    print("🔍 شروع پردازش...")
    print(f"📁 مسیر: {folder_path}")
    print("=" * 60)
    
    try:
        import openpyxl
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    except ImportError:
        print("❌ خطا: کتابخانه openpyxl نصب نیست!")
        print("💡 برای نصب از دستور زیر استفاده کنید:")
        print("   pip install openpyxl")
        return
    
    all_files = []
    extensions = ['*.pdf', '*.doc', '*.docx', '*.xls', '*.xlsx']
    
    for ext in extensions:
        all_files.extend(list(Path(folder_path).glob(ext)))
    
    if not all_files:
        print("❌ هیچ فایلی پیدا نشد!")
        return
    
    print(f"\n✅ تعداد {len(all_files)} فایل پیدا شد:")
    for f in all_files:
        print(f"   📄 {f.name}")
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Weekly Reports"
    
    ws.column_dimensions['A'].width = 50
    ws.column_dimensions['B'].width = 55
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 40
    
    headers = ['Report No', 'File Name (Link)', 'Date', 'Report Title']
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col)
        cell.value = header
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = thin_border
    
    print("\n🔄 در حال پردازش فایل‌ها...")
    
    for idx, file in enumerate(sorted(all_files), 2):
        filename = file.name
        file_path = str(file.absolute())
        file_ext = file.suffix.lower()
        
        print(f"   ⏳ پردازش: {filename}")
        
        report_no = extract_report_no_from_filename(filename)
        
        if file_ext == '.pdf':
            date = extract_date_from_pdf(file_path)
            report_title = extract_report_title_from_pdf(file_path)
        elif file_ext in ['.doc', '.docx']:
            date = extract_date_from_word(file_path)
            report_title = extract_report_title_from_word(file_path)
        else:
            date = "N/A"
            report_title = "N/A"
        
        cell_a = ws.cell(row=idx, column=1)
        cell_a.value = report_no
        cell_a.alignment = Alignment(horizontal='left', vertical='center')
        cell_a.border = thin_border
        
        cell_b = ws.cell(row=idx, column=2)
        cell_b.value = filename
        cell_b.hyperlink = file_path
        cell_b.font = Font(color="0563C1", underline="single")
        cell_b.alignment = Alignment(horizontal='left', vertical='center')
        cell_b.border = thin_border
        
        cell_c = ws.cell(row=idx, column=3)
        cell_c.value = date
        cell_c.alignment = Alignment(horizontal='center', vertical='center')
        cell_c.border = thin_border
        
        cell_d = ws.cell(row=idx, column=4)
        cell_d.value = report_title
        cell_d.alignment = Alignment(horizontal='center', vertical='center')
        cell_d.border = thin_border
    
    try:
        output_path = os.path.join(folder_path, output_file)
        wb.save(output_path)
        print("\n" + "=" * 60)
        print(f"✅ فایل اکسل با موفقیت ایجاد شد!")
        print(f"📂 مسیر فایل: {output_path}")
        print(f"📊 تعداد فایل‌های پردازش شده: {len(all_files)}")
        print("=" * 60)
    except PermissionError:
        print(f"\n❌ خطا: دسترسی به ذخیره فایل وجود ندارد!")
        print(f"💡 لطفاً مطمئن شوید فایل {output_file} باز نیست.")
    except Exception as e:
        print(f"\n❌ خطا در ذخیره فایل: {e}")

if __name__ == "__main__":
    folder_path = r"D:\Sepher_Pasargad\works\Production\Quarterly"
    
    if not os.path.exists(folder_path):
        print(f"❌ مسیر {folder_path} وجود ندارد!")
    else:
        create_excel_report(folder_path)