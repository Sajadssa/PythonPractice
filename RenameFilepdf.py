#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Maintenance Monthly Reports - File Renamer
نسخه بهبود یافته با رفع مشکلات

این اسکریپت فایل‌های PDF و Word گزارشات نگهداری ماهانه را
به صورت خودکار با فرمت استاندارد تغییر نام می‌دهد.
"""

import os
from pathlib import Path
import re
from datetime import datetime
import sys
from collections import defaultdict

# کتابخانه‌های اصلی
try:
    import PyPDF2
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from docx import Document
except ImportError as e:
    print(f"❌ خطا: کتابخانه موردنیاز نصب نیست: {e}")
    print("لطفا با دستور زیر کتابخانه‌ها را نصب کنید:")
    print("pip install PyPDF2 openpyxl python-docx")
    sys.exit(1)

# کتابخانه‌های OCR (اختیاری)
try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("⚠️ توجه: کتابخانه‌های OCR نصب نیستند. برای PDF های اسکن شده از OCR استفاده نخواهد شد.")


def parse_date_to_excel(date_str):
    """
    تبدیل تاریخ به datetime object
    ورودی: 14-Oct-2024 یا 14-October-2024 یا 14 Oct 2024
    """
    if not date_str:
        return None
    
    try:
        months = {
            'jan': 1, 'january': 1,
            'feb': 2, 'february': 2,
            'mar': 3, 'march': 3,
            'apr': 4, 'april': 4,
            'may': 5,
            'jun': 6, 'june': 6,
            'jul': 7, 'july': 7,
            'aug': 8, 'august': 8,
            'sep': 9, 'september': 9,
            'oct': 10, 'october': 10,
            'nov': 11, 'november': 11,
            'dec': 12, 'december': 12
        }
        
        # پاک کردن فضاهای اضافی و تبدیل به lowercase
        date_str = date_str.strip().lower()
        
        # فرمت 1: 14-Oct-2024 یا 14-october-2024
        if '-' in date_str:
            parts = date_str.split('-')
            if len(parts) == 3:
                day = int(parts[0])
                month_name = parts[1]
                year = int(parts[2])
                
                month = months.get(month_name)
                if month:
                    return datetime(year, month, day)
        
        # فرمت 2: 14 Oct 2024
        parts = date_str.split()
        if len(parts) == 3:
            day = int(parts[0])
            month_name = parts[1]
            year = int(parts[2])
            
            month = months.get(month_name)
            if month:
                return datetime(year, month, day)
                
    except Exception as e:
        print(f"   ⚠️ خطا در تبدیل تاریخ '{date_str}': {e}")
    
    return None


def extract_text_from_pdf_with_ocr(pdf_path):
    """
    استخراج متن از PDF با OCR (اختیاری)
    """
    if not OCR_AVAILABLE:
        return ""
    
    try:
        print(f"   🔍 تلاش برای OCR...")
        images = convert_from_path(pdf_path, first_page=1, last_page=1, dpi=300)
        
        if images:
            text = pytesseract.image_to_string(images[0], lang='eng')
            return text
    except Exception as e:
        print(f"   ⚠️ خطا در OCR: {str(e)}")
    
    return ""


def extract_info_from_pdf(pdf_path):
    """
    استخراج اطلاعات از PDF:
    - Document No (مانند: SJSC-GGNRSP-MADR-REMO-2024-G01)
    - Date (از جدول دوم)
    - Report Title
    - Period (From ... to ...)
    """
    text = ""
    
    try:
        # خواندن مستقیم PDF
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            if len(pdf_reader.pages) > 0:
                first_page = pdf_reader.pages[0]
                text = first_page.extract_text()
        
        # اگر متن کافی نبود، از OCR استفاده کن (اگر موجود باشد)
        if (not text or len(text.strip()) < 50) and OCR_AVAILABLE:
            print(f"   ⚠️ PDF اسکن شده شناسایی شد، استفاده از OCR...")
            text = extract_text_from_pdf_with_ocr(pdf_path)
        
        if text:
            print(f"   📄 متن استخراج شده ({len(text)} کاراکتر)")
            
            # استخراج Document No
            # الگوهای مختلف برای Document No
            patterns_docno = [
                # فرمت استاندارد: SJSC-XXX-XXX-REMO-YYYY-GNN
                r'Document\s*No\.?\s*:?\s*(SJSC-[A-Z0-9]+-[A-Z0-9]+-[A-Z]+-(\d{4})-(G\d{2}))',
                r'Document\s*Number\s*:?\s*(SJSC-[A-Z0-9]+-[A-Z0-9]+-[A-Z]+-(\d{4})-(G\d{2}))',
                r'Doc\s*No\.?\s*:?\s*(SJSC-[A-Z0-9]+-[A-Z0-9]+-[A-Z]+-(\d{4})-(G\d{2}))',
                # فرمت کلی
                r'(SJSC-[A-Z0-9]+-[A-Z0-9]+-REMO-(\d{4})-(G\d{2}))',
                r'(SJSC-[A-Z0-9]+-[A-Z0-9]+-[A-Z]+-(\d{4})-(G\d{2}))',
            ]
            
            doc_no = None
            doc_number = None
            rev = None
            
            for pattern in patterns_docno:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    groups = match.groups()
                    if len(groups) >= 3:
                        doc_no = groups[0]
                        doc_number = groups[1]
                        rev = groups[2]
                        print(f"   ✅ Document No: {doc_no}")
                        break
            
            # استخراج Date از جدول دوم
            patterns_date = [
                r'Date\s*:?\s*([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})',
                r'Date\s+([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})',
                r'Approved\s+by\s+Date\s+([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})',
                # الگوهای اضافی
                r'Date:\s*([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})',
            ]
            
            date_obj = None
            date_str = None
            
            for pattern in patterns_date:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    date_str = match.group(1).replace(' ', '-')
                    date_obj = parse_date_to_excel(date_str)
                    if date_obj:
                        print(f"   ✅ Date: {date_str} → {date_obj.strftime('%d/%m/%Y')}")
                        break
                if date_obj:
                    break
            
            # استخراج عنوان گزارش
            report_title = None
            text_upper = text.upper()
            if 'MAINTENANCE' in text_upper and 'MONTHLY' in text_upper and 'REPORT' in text_upper:
                report_title = 'MAINTENANCE MONTHLY REPORT'
            
            # استخراج دوره گزارش (From ... to ...)
            period = None
            period_patterns = [
                r'\(From\s+([0-9]{1,2}[-\s][A-Za-z]{3,9})\s+to\s+([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})\)',
                r'From\s+([0-9]{1,2}[-\s][A-Za-z]{3,9})\s+to\s+([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})',
            ]
            
            for pattern in period_patterns:
                period_match = re.search(pattern, text, re.IGNORECASE)
                if period_match:
                    period = f"From {period_match.group(1)} to {period_match.group(2)}"
                    print(f"   ✅ Period: {period}")
                    break
            
            return {
                'doc_no': doc_no,
                'doc_number': doc_number,
                'rev': rev,
                'date': date_obj,
                'date_str': date_str,
                'report_title': report_title,
                'period': period
            }
                    
    except Exception as e:
        print(f"   ⚠️ خطا در خواندن PDF: {str(e)}")
        import traceback
        traceback.print_exc()
    
    return None


def extract_info_from_word(word_path):
    """
    استخراج اطلاعات از فایل Word (.docx)
    """
    try:
        doc = Document(word_path)
        
        # متن کامل سند
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # متن جداول
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        text = '\n'.join(full_text)
        print(f"   📄 متن Word استخراج شده ({len(text)} کاراکتر)")
        
        # استخراج Document No
        patterns_docno = [
            r'Document\s*No\.?\s*:?\s*(SJSC-[A-Z0-9]+-[A-Z0-9]+-[A-Z]+-(\d{4})-(G\d{2}))',
            r'Document\s*Number\s*:?\s*(SJSC-[A-Z0-9]+-[A-Z0-9]+-[A-Z]+-(\d{4})-(G\d{2}))',
            r'(SJSC-[A-Z0-9]+-[A-Z0-9]+-REMO-(\d{4})-(G\d{2}))',
            r'(SJSC-[A-Z0-9]+-[A-Z0-9]+-[A-Z]+-(\d{4})-(G\d{2}))',
        ]
        
        doc_no = None
        doc_number = None
        rev = None
        
        for pattern in patterns_docno:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                groups = match.groups()
                if len(groups) >= 3:
                    doc_no = groups[0]
                    doc_number = groups[1]
                    rev = groups[2]
                    print(f"   ✅ Document No: {doc_no}")
                    break
        
        # استخراج Date
        patterns_date = [
            r'Date\s*:?\s*([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})',
            r'Date\s+([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})',
            r'Date:\s*([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})',
        ]
        
        date_obj = None
        date_str = None
        
        for pattern in patterns_date:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                date_str = match.group(1).replace(' ', '-')
                date_obj = parse_date_to_excel(date_str)
                if date_obj:
                    print(f"   ✅ Date: {date_str} → {date_obj.strftime('%d/%m/%Y')}")
                    break
            if date_obj:
                break
        
        # عنوان گزارش
        report_title = None
        text_upper = text.upper()
        if 'MAINTENANCE' in text_upper and 'MONTHLY' in text_upper and 'REPORT' in text_upper:
            report_title = 'MAINTENANCE MONTHLY REPORT'
        
        # دوره گزارش
        period = None
        period_patterns = [
            r'\(From\s+([0-9]{1,2}[-\s][A-Za-z]{3,9})\s+to\s+([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})\)',
            r'From\s+([0-9]{1,2}[-\s][A-Za-z]{3,9})\s+to\s+([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})',
        ]
        
        for pattern in period_patterns:
            period_match = re.search(pattern, text, re.IGNORECASE)
            if period_match:
                period = f"From {period_match.group(1)} to {period_match.group(2)}"
                print(f"   ✅ Period: {period}")
                break
        
        return {
            'doc_no': doc_no,
            'doc_number': doc_number,
            'rev': rev,
            'date': date_obj,
            'date_str': date_str,
            'report_title': report_title,
            'period': period
        }
    
    except Exception as e:
        print(f"   ⚠️ خطا در خواندن Word: {str(e)}")
        import traceback
        traceback.print_exc()
    
    return None


def create_excel_report(files_data, output_path):
    """
    ایجاد گزارش اکسل با فرمت زیبا
    """
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Maintenance Reports"
        
        # استایل‌های زیبا
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # هدرها
        headers = ['ردیف', 'نام فایل اصلی', 'نام فایل جدید', 'عنوان گزارش', 
                   'دوره', 'Document No', 'شماره', 'REV', 'تاریخ', 'وضعیت']
        ws.append(headers)
        
        # استایل هدر
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border
        
        # داده‌ها
        for idx, data in enumerate(files_data, start=1):
            date_value = data['date'] if data['date'] else 'N/A'
            
            row = [
                idx,
                data['old_name'],
                data['new_name'] if data['new_name'] else 'N/A',
                data['report_title'] if data['report_title'] else 'N/A',
                data['period'] if data['period'] else 'N/A',
                data['doc_no'] if data['doc_no'] else 'N/A',
                data['doc_number'] if data['doc_number'] else 'N/A',
                data['rev'] if data['rev'] else 'N/A',
                date_value,
                data['status']
            ]
            ws.append(row)
            
            # استایل ردیف
            row_num = idx + 1
            for col_idx, cell in enumerate(ws[row_num], start=1):
                cell.border = border
                cell.alignment = Alignment(horizontal='center', vertical='center')
                
                # فرمت تاریخ
                if col_idx == 9 and isinstance(cell.value, datetime):
                    cell.number_format = 'DD/MM/YYYY'
        
        # تنظیم عرض ستون‌ها
        column_widths = [8, 40, 45, 35, 30, 40, 12, 8, 15, 20]
        for idx, width in enumerate(column_widths, start=1):
            ws.column_dimensions[chr(64 + idx)].width = width
        
        # ذخیره فایل
        wb.save(output_path)
        print(f"\n📊 فایل اکسل ایجاد شد: {output_path}")
        
    except Exception as e:
        print(f"❌ خطا در ایجاد فایل اکسل: {str(e)}")
        import traceback
        traceback.print_exc()


def rename_files(folder_path, dry_run=False):
    """
    تغییر نام فایل‌های PDF و Word
    
    Args:
        folder_path: مسیر پوشه حاوی فایل‌ها
        dry_run: اگر True باشد، فقط شبیه‌سازی می‌کند و فایل‌ها را تغییر نام نمی‌دهد
    """
    print("="*80)
    print("🔄 تغییر نام Maintenance Monthly Reports")
    print("="*80)
    print(f"📂 مسیر پوشه: {folder_path}")
    if dry_run:
        print("🔍 حالت تست (Dry Run) - فایل‌ها تغییر نام نخواهند خورد")
    print()
    
    # پیدا کردن فایل‌ها
    pdf_files = list(Path(folder_path).glob('*.pdf'))
    word_files = list(Path(folder_path).glob('*.docx')) + list(Path(folder_path).glob('*.doc'))
    
    # حذف فایل‌هایی که قبلاً با فرمت استاندارد نام‌گذاری شده‌اند
    pdf_files = [f for f in pdf_files if not f.name.startswith('SJSC-GGNRSP-MADR-REMO-')]
    word_files = [f for f in word_files if not f.name.startswith('SJSC-GGNRSP-MADR-REMO-')]
    
    all_files = pdf_files + word_files
    
    if not all_files:
        print("❌ هیچ فایلی برای پردازش پیدا نشد!")
        print("   (فایل‌های با فرمت SJSC-GGNRSP-MADR-REMO-* از لیست حذف شدند)")
        return
    
    print(f"📁 {len(pdf_files)} فایل PDF و {len(word_files)} فایل Word پیدا شد\n")
    print("🔍 در حال استخراج اطلاعات...")
    print("-"*80)
    
    # استخراج اطلاعات
    files_data = []
    
    for file_path in all_files:
        print(f"\n📄 پردازش: {file_path.name}")
        
        try:
            if file_path.suffix.lower() == '.pdf':
                info = extract_info_from_pdf(file_path)
            else:
                info = extract_info_from_word(file_path)
            
            if info and info['doc_number'] and info['rev']:
                files_data.append({
                    'path': file_path,
                    'old_name': file_path.name,
                    'doc_no': info['doc_no'],
                    'doc_number': info['doc_number'],
                    'rev': info['rev'],
                    'date': info['date'],
                    'date_str': info['date_str'],
                    'report_title': info['report_title'],
                    'period': info['period'],
                    'new_name': None,
                    'status': 'در انتظار'
                })
            else:
                files_data.append({
                    'path': file_path,
                    'old_name': file_path.name,
                    'doc_no': None,
                    'doc_number': None,
                    'rev': None,
                    'date': None,
                    'date_str': None,
                    'report_title': None,
                    'period': None,
                    'new_name': None,
                    'status': 'خطا - اطلاعات کافی یافت نشد'
                })
                print(f"   ❌ نتوانستیم اطلاعات لازم را استخراج کنیم!")
        
        except Exception as e:
            print(f"   ❌ خطای غیرمنتظره: {str(e)}")
            files_data.append({
                'path': file_path,
                'old_name': file_path.name,
                'doc_no': None,
                'doc_number': None,
                'rev': None,
                'date': None,
                'date_str': None,
                'report_title': None,
                'period': None,
                'new_name': None,
                'status': f'خطا: {str(e)}'
            })
    
    print("-"*80)
    
    # شناسایی تکراری‌ها و تخصیص نام جدید
    print("\n🔢 بررسی تکراری‌ها و تخصیص نام...")
    
    groups = defaultdict(list)
    for data in files_data:
        if data['doc_number'] and data['rev']:
            key = f"{data['doc_number']}-{data['rev']}"
            groups[key].append(data)
    
    # تخصیص نام جدید
    for key, group in groups.items():
        if len(group) == 1:
            data = group[0]
            ext = data['path'].suffix
            data['new_name'] = f"SJSC-GGNRSP-MADR-REMO-{data['doc_number']}-{data['rev']}{ext}"
        else:
            print(f"   ⚠️ فایل تکراری یافت شد: {key}")
            for idx, data in enumerate(group, start=1):
                ext = data['path'].suffix
                data['new_name'] = f"SJSC-GGNRSP-MADR-REMO-{data['doc_number']}-{data['rev']}_copy{idx}{ext}"
                print(f"      → نام فایل {idx}: ...{data['new_name']}")
    
    # تغییر نام فایل‌ها
    if not dry_run:
        print("\n🔄 شروع تغییر نام فایل‌ها...")
        print("-"*80)
        
        renamed_count = 0
        failed_count = 0
        
        for data in files_data:
            if data['new_name']:
                old_path = data['path']
                new_name = data['new_name']
                new_path = old_path.parent / new_name
                
                # بررسی تکراری بودن نام
                if new_path.exists() and new_path != old_path:
                    print(f"⚠️ فایل با این نام وجود دارد: {new_name}")
                    data['status'] = 'رد شده - نام تکراری در سیستم فایل'
                    failed_count += 1
                    continue
                
                try:
                    old_path.rename(new_path)
                    renamed_count += 1
                    data['status'] = '✅ موفق'
                    print(f"✅ {old_path.name}")
                    print(f"   ➜ {new_name}")
                except Exception as e:
                    print(f"❌ خطا در تغییر نام: {str(e)}")
                    data['status'] = f'❌ خطا: {str(e)}'
                    failed_count += 1
            else:
                failed_count += 1
        
        print("-"*80)
    else:
        print("\n🔍 پیش‌نمایش تغییرات (Dry Run):")
        print("-"*80)
        
        renamed_count = 0
        failed_count = 0
        
        for data in files_data:
            if data['new_name']:
                print(f"✅ {data['old_name']}")
                print(f"   ➜ {data['new_name']}")
                renamed_count += 1
                data['status'] = '🔍 شبیه‌سازی موفق'
            else:
                failed_count += 1
        
        print("-"*80)
    
    # ایجاد گزارش اکسل
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    excel_filename = f"Maintenance_Rename_Report_{timestamp}.xlsx"
    excel_path = Path(folder_path) / excel_filename
    
    create_excel_report(files_data, excel_path)
    
    # خلاصه نهایی
    print(f"\n📊 خلاصه نتایج:")
    print(f"   ✅ موفق: {renamed_count}")
    print(f"   ❌ ناموفق: {failed_count}")
    print(f"   📝 کل فایل‌ها: {len(files_data)}")
    print("="*80)
    
    return excel_path


def main():
    """
    تابع اصلی برنامه
    """
    # تنظیم مسیر پوشه
    # در ویندوز:
    FOLDER_PATH = r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\monthly"
    
    # در لینوکس/مک:
    # FOLDER_PATH = "/path/to/your/folder"
    
    # استفاده از مسیر فعلی برای تست
    FOLDER_PATH = os.path.join(os.getcwd(), "test_reports")
    
    print("\n" + "="*80)
    print("🔧 Maintenance Monthly Reports - File Renamer")
    print("نسخه بهبود یافته")
    print("="*80)
    
    # بررسی وجود پوشه
    if not os.path.exists(FOLDER_PATH):
        print(f"\n❌ خطا: پوشه پیدا نشد!")
        print(f"مسیر: {FOLDER_PATH}\n")
        
        # پیشنهاد ایجاد پوشه تست
        print("💡 برای تست، یک پوشه نمونه ایجاد می‌کنیم...")
        try:
            os.makedirs(FOLDER_PATH, exist_ok=True)
            print(f"✅ پوشه تست ایجاد شد: {FOLDER_PATH}")
            print("\nلطفا فایل‌های PDF یا Word خود را در این پوشه قرار دهید و دوباره اجرا کنید.")
        except Exception as e:
            print(f"❌ خطا در ایجاد پوشه: {e}")
        
        return
    
    # دریافت تایید کاربر
    print(f"\n⚠️ هشدار: این برنامه نام فایل‌های PDF و Word را تغییر می‌دهد!")
    print(f"📂 پوشه: {FOLDER_PATH}")
    print("\nگزینه‌ها:")
    print("  1. اجرای واقعی (تغییر نام فایل‌ها)")
    print("  2. حالت تست (فقط پیش‌نمایش بدون تغییر)")
    print("  3. لغو")
    
    choice = input("\nانتخاب شما (1/2/3): ").strip()
    
    if choice == '1':
        print("\n✅ اجرای واقعی شروع می‌شود...\n")
        excel_path = rename_files(FOLDER_PATH, dry_run=False)
    elif choice == '2':
        print("\n🔍 حالت تست (Dry Run) شروع می‌شود...\n")
        excel_path = rename_files(FOLDER_PATH, dry_run=True)
    else:
        print("\n❌ عملیات لغو شد.")
        return
    
    print(f"\n✨ کار تمام شد!")
    if excel_path:
        print(f"📊 گزارش اکسل: {excel_path}")


if __name__ == "__main__":
    main()