import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from deep_translator import GoogleTranslator
import time
import re
import os

class OilGasTranslator:
    """
    مترجم تخصصی اسناد نفت و گاز
    """
    
    def __init__(self):
        # دیکشنری اصطلاحات تخصصی نفت و گاز
        self.technical_terms = {
            'Maintenance': 'نگهداری و تعمیرات',
            'Contractor': 'پیمانکار',
            'Company': 'کارفرما',
            'CMMS': 'سیستم مدیریت نگهداری کامپیوتری',
            'ESD': 'خاموش‌سازی اضطراری',
            'Emergency Shut Down': 'خاموش‌سازی اضطراری',
            'HVAC': 'تهویه مطبوع',
            'Heat, Ventilation and Air Conditioning': 'گرمایش، تهویه و تهویه مطبوع',
            'Wellhead': 'دهانه چاه',
            'Well Head Control Panel': 'پنل کنترل دهانه چاه',
            'WHCP': 'پنل کنترل دهانه چاه',
            'PPE': 'تجهیزات حفاظت فردی',
            'Personal Protective Equipment': 'تجهیزات حفاظت فردی',
            'PM': 'نگهداری پیشگیرانه',
            'Preventive Maintenance': 'نگهداری پیشگیرانه',
            'CM': 'نگهداری اصلاحی',
            'Corrective Maintenance': 'نگهداری اصلاحی',
            'LBV': 'شیر قطع خط',
            'Line Break Valve': 'شیر قطع خط',
            'Receiving Area': 'منطقه دریافت',
            'RA': 'منطقه دریافت',
            'Booster Cluster': 'خوشه تقویت فشار',
            'Separator': 'جداساز',
            'Flow Line': 'خط جریان',
            'Pipeline': 'خط لوله',
            'Valve': 'شیر',
            'Actuator': 'محرک',
            'Safety Valve': 'شیر اطمینان',
            'Relief Valve': 'شیر تخلیه فشار',
            'Diesel Generator': 'دیزل ژنراتور',
            'HSE': 'بهداشت، ایمنی و محیط زیست',
            'Health, Safety and Environment': 'بهداشت، ایمنی و محیط زیست',
            'Calibration': 'کالیبراسیون',
            'Inspection': 'بازرسی',
            'Overhaul': 'بازسازی کامل',
            'Corrosion': 'خوردگی',
            'Pressure Vessel': 'مخزن تحت فشار',
            'PSV': 'شیر اطمینان',
            'ESDV': 'شیر خاموش‌سازی اضطراری',
            'HIPPS': 'سیستم حفاظتی با فشار بالا',
            'MOV': 'شیر موتوردار',
            'DCS': 'سیستم کنترل توزیع شده',
            'PLC': 'کنترلر منطقی برنامه‌پذیر',
            'SCADA': 'سیستم نظارت و کنترل',
            'Instrumentation': 'ابزار دقیق',
            'Transmitter': 'ترنسمیتر',
            'X-mas Tree': 'درخت کریسمس (مجموعه شیرهای دهانه چاه)',
            'Manifold': 'منیفولد',
            'Gathering': 'جمع‌آوری',
            'Production': 'تولید',
            'Facility': 'تاسیسات',
            'Equipment': 'تجهیزات',
            'Spare Parts': 'قطعات یدکی',
            'Consumables': 'مواد مصرفی',
            'Work Order': 'دستور کار',
            'Scope of Work': 'شرح خدمات',
            'Contract': 'قرارداد',
            'Mobilization': 'بسیج منابع',
            'Demobilization': 'جمع‌آوری منابع',
            'TPI': 'بازرس شخص ثالث',
            'Third Party Inspector': 'بازرس شخص ثالث',
            'ITP': 'برنامه بازرسی و آزمون',
            'Inspection and Test Plan': 'برنامه بازرسی و آزمون',
            'Chemical Injection': 'تزریق مواد شیمیایی',
            'Pump': 'پمپ',
            'Compressor': 'کمپرسور',
            'Transformer': 'ترانسفورماتور',
            'Switchgear': 'تابلو برق',
            'UPS': 'منبع تغذیه بدون وقفه',
            'Fire Fighting': 'اطفاء حریق',
            'Piping': 'لوله‌کشی',
            'Welding': 'جوشکاری',
            'Hydro Test': 'آزمون هیدرواستاتیک',
            'Leak Test': 'آزمون نشتی',
            'Non-Destructive Test': 'آزمون غیرمخرب',
            'NDT': 'آزمون غیرمخرب',
            'Sandblasting': 'سندبلاست',
            'Painting': 'رنگ‌آمیزی',
            'Coating': 'پوشش‌دهی',
            'Cathodic Protection': 'حفاظت کاتدی',
            'Civil Work': 'عملیات عمرانی',
            'Road': 'جاده',
            'Foundation': 'فونداسیون',
            'Concrete': 'بتن',
            'Gravel': 'شن',
            'Sand': 'ماسه',
            'Cement': 'سیمان',
            'Rebar': 'میلگرد',
            'Excavation': 'خاکبرداری',
            'Compaction': 'تراکم',
            'Scaffolding': 'داربست',
            'Crane': 'جرثقیل',
            'Forklift': 'لیفتراک',
            'Boom Truck': 'کامیون بوم‌دار',
            'Personnel': 'پرسنل',
            'Technician': 'تکنسین',
            'Engineer': 'مهندس',
            'Supervisor': 'سرپرست',
            'Coordinator': 'هماهنگ‌کننده',
            'Planner': 'برنامه‌ریز',
            'Camp': 'کمپ',
            'Accommodation': 'محل اسکان',
            'Food': 'غذا',
            'Catering': 'تغذیه',
            'Transportation': 'حمل و نقل',
            'Security': 'امنیت',
            'Training': 'آموزش',
        }
    
    def translate_with_context(self, text, chunk_size=4500):
        """
        ترجمه متن با در نظر گرفتن اصطلاحات تخصصی
        """
        if not text or len(text.strip()) == 0:
            return text
        
        # جایگزینی اصطلاحات تخصصی قبل از ترجمه
        original_text = text
        temp_replacements = {}
        
        for idx, (en_term, fa_term) in enumerate(self.technical_terms.items()):
            placeholder = f"__TECH_TERM_{idx}__"
            if en_term in text:
                text = text.replace(en_term, placeholder)
                temp_replacements[placeholder] = fa_term
        
        try:
            # ترجمه متن
            if len(text) > chunk_size:
                # تقسیم به بخش‌های کوچکتر
                sentences = text.split('.')
                translated_parts = []
                current_chunk = ""
                
                for sentence in sentences:
                    if len(current_chunk) + len(sentence) < chunk_size:
                        current_chunk += sentence + "."
                    else:
                        if current_chunk:
                            trans = GoogleTranslator(source='en', target='fa').translate(current_chunk)
                            translated_parts.append(trans)
                            time.sleep(0.5)
                        current_chunk = sentence + "."
                
                if current_chunk:
                    trans = GoogleTranslator(source='en', target='fa').translate(current_chunk)
                    translated_parts.append(trans)
                
                translated_text = " ".join(translated_parts)
            else:
                translated_text = GoogleTranslator(source='en', target='fa').translate(text)
            
            # بازگرداندن اصطلاحات تخصصی
            for placeholder, fa_term in temp_replacements.items():
                translated_text = translated_text.replace(placeholder, fa_term)
            
            return translated_text
        
        except Exception as e:
            print(f"خطا در ترجمه: {e}")
            return original_text
    
    def extract_pdf_structure(self, pdf_path):
        """
        استخراج ساختار دقیق PDF شامل متن، فونت، اندازه، جداول و...
        """
        doc = fitz.open(pdf_path)
        pages_content = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_dict = page.get_text("dict")
            
            page_structure = {
                'page_number': page_num + 1,
                'blocks': []
            }
            
            for block in page_dict["blocks"]:
                if block["type"] == 0:  # Text block
                    block_info = {
                        'type': 'text',
                        'bbox': block["bbox"],
                        'lines': []
                    }
                    
                    for line in block["lines"]:
                        line_text = ""
                        font_size = 0
                        is_bold = False
                        
                        for span in line["spans"]:
                            line_text += span["text"]
                            font_size = max(font_size, span["size"])
                            if "bold" in span["font"].lower():
                                is_bold = True
                        
                        block_info['lines'].append({
                            'text': line_text.strip(),
                            'font_size': font_size,
                            'is_bold': is_bold
                        })
                    
                    page_structure['blocks'].append(block_info)
            
            pages_content.append(page_structure)
        
        doc.close()
        return pages_content
    
    def set_rtl_paragraph(self, paragraph):
        """
        تنظیم جهت راست به چپ برای پاراگراف
        """
        pPr = paragraph._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
    
    def create_bilingual_document(self, pdf_path, output_path):
        """
        ایجاد سند دو زبانه: صفحه فرد انگلیسی، صفحه زوج فارسی
        """
        print("=" * 70)
        print("شروع فرآیند ترجمه تخصصی سند نفت و گاز")
        print("=" * 70)
        
        # استخراج ساختار PDF
        print("\n[1/4] در حال استخراج ساختار PDF...")
        pages_structure = self.extract_pdf_structure(pdf_path)
        print(f"✓ تعداد صفحات: {len(pages_structure)}")
        
        # ایجاد سند Word
        doc = Document()
        
        # تنظیمات فونت فارسی
        style = doc.styles['Normal']
        font = style.font
        font.name = 'B Nazanin'
        font.size = Pt(11)
        
        print("\n[2/4] شروع ترجمه صفحات...")
        
        for page_idx, page_structure in enumerate(pages_structure):
            page_num = page_structure['page_number']
            print(f"\n--- صفحه {page_num} از {len(pages_structure)} ---")
            
            # صفحه انگلیسی (متن اصلی)
            print(f"  • افزودن متن اصلی انگلیسی...")
            
            # عنوان صفحه
            header = doc.add_paragraph()
            header_run = header.add_run(f"Page {page_num} - Original Text (English)")
            header_run.bold = True
            header_run.font.size = Pt(14)
            header_run.font.color.rgb = RGBColor(0, 51, 102)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # فاصله
            
            # افزودن محتوای انگلیسی
            for block in page_structure['blocks']:
                for line in block['lines']:
                    if line['text']:
                        p = doc.add_paragraph()
                        run = p.add_run(line['text'])
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(line['font_size'] if line['font_size'] > 0 else 11)
                        run.bold = line['is_bold']
                        
                        # تشخیص عنوان
                        if line['is_bold'] or line['font_size'] > 12:
                            run.font.color.rgb = RGBColor(0, 0, 0)
            
            # شکست صفحه
            doc.add_page_break()
            
            # صفحه فارسی (ترجمه)
            print(f"  • ترجمه به فارسی...")
            
            # عنوان صفحه فارسی
            header_fa = doc.add_paragraph()
            header_fa_run = header_fa.add_run(f"صفحه {page_num} - ترجمه تخصصی (فارسی)")
            header_fa_run.bold = True
            header_fa_run.font.size = Pt(14)
            header_fa_run.font.name = 'B Nazanin'
            header_fa_run.font.color.rgb = RGBColor(0, 51, 102)
            header_fa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.set_rtl_paragraph(header_fa)
            
            doc.add_paragraph()  # فاصله
            
            # ترجمه و افزودن محتوای فارسی
            block_count = len(page_structure['blocks'])
            for block_idx, block in enumerate(page_structure['blocks']):
                for line_idx, line in enumerate(block['lines']):
                    if line['text']:
                        # ترجمه
                        translated = self.translate_with_context(line['text'])
                        
                        if translated:
                            p = doc.add_paragraph()
                            run = p.add_run(translated)
                            run.font.name = 'B Nazanin'
                            run.font.size = Pt(line['font_size'] if line['font_size'] > 0 else 11)
                            run.bold = line['is_bold']
                            
                            # راست‌چین کردن
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            self.set_rtl_paragraph(p)
                        
                        # نمایش پیشرفت
                        progress = ((block_idx + 1) / block_count) * 100
                        print(f"    پیشرفت: {progress:.0f}%", end='\r')
                        
                        time.sleep(0.3)  # تاخیر برای API
            
            print(f"    پیشرفت: 100% ✓")
            
            # شکست صفحه برای صفحه بعدی (مگر آخرین صفحه)
            if page_idx < len(pages_structure) - 1:
                doc.add_page_break()
        
        # ذخیره سند
        print(f"\n[3/4] در حال ذخیره سند...")
        doc.save(output_path)
        print(f"✓ فایل ذخیره شد: {output_path}")
        
        print("\n[4/4] خلاصه کار:")
        print(f"  • تعداد صفحات ترجمه شده: {len(pages_structure)}")
        print(f"  • تعداد صفحات خروجی: {len(pages_structure) * 2}")
        print(f"  • فرمت: صفحات فرد = انگلیسی، صفحات زوج = فارسی")
        
        print("\n" + "=" * 70)
        print("✓✓✓ ترجمه با موفقیت کامل شد! ✓✓✓")
        print("=" * 70)

# اجرای برنامه
if __name__ == "__main__":
    # مسیر فایل PDF
    pdf_path = r"D:\Sepher_Pasargad\works\Maintenace\PythonDataAnalysis\PythonPractice\2. Appendix No. 1- Scope of Work Rev 14.pdf"
    
    # مسیر خروجی
    output_path = r"D:\Sepher_Pasargad\works\Maintenace\PythonDataAnalysis\PythonPractice\Translated_Maintenance_Contract_FA.docx"
    
    # بررسی وجود فایل
    if not os.path.exists(pdf_path):
        print(f"خطا: فایل PDF در مسیر زیر یافت نشد:")
        print(pdf_path)
    else:
        # ایجاد مترجم و شروع ترجمه
        translator = OilGasTranslator()
        translator.create_bilingual_document(pdf_path, output_path)
        
        print(f"\n✓ فایل خروجی در مسیر زیر ذخیره شد:")
        print(output_path)