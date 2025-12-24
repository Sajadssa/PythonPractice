#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import re
from pathlib import Path
from docx import Document

def test_one_file(docx_path):
    """Test extraction on ONE file with full debug"""
    print("="*70)
    print(f"Testing: {docx_path.name}")
    print("="*70 + "\n")
    
    doc = Document(docx_path)
    
    if not doc.tables:
        print("NO TABLES!")
        return
    
    print(f"Total tables: {len(doc.tables)}\n")
    
    # CHECK ALL TABLES to find the one with "Sequence Number"
    header_table_idx = None
    
    for table_idx, table in enumerate(doc.tables):
        print(f"TABLE {table_idx}:")
        print("-"*70)
        
        # Check first row
        if len(table.rows) > 0:
            first_row = [c.text.strip() for c in table.rows[0].cells]
            print(f"First row: {first_row}")
            
            # Check if this has "Sequence Number"
            for cell_text in first_row:
                if 'Sequence' in cell_text and 'Number' in cell_text:
                    header_table_idx = table_idx
                    print(f"✓✓✓ FOUND HEADER TABLE! ✓✓✓")
                    break
        
        print()
        
        if header_table_idx is not None:
            break
    
    if header_table_idx is None:
        print("❌ NO TABLE WITH 'Sequence Number' FOUND!")
        return
    
    # Now extract from the correct table
    print("="*70)
    print(f"EXTRACTING FROM TABLE {header_table_idx}:")
    print("="*70 + "\n")
    
    header_table = doc.tables[header_table_idx]
    
    # Show all rows
    for row_idx, row in enumerate(header_table.rows):
        cells = [c.text.strip() for c in row.cells]
        print(f"Row {row_idx}: {cells}")
    
    print("\nEXTRACTION:")
    print("-"*70)
    
    # Get headers (row 0)
    headers = [c.text.strip() for c in header_table.rows[0].cells]
    print(f"Headers: {headers}")
    
    # Get values (row 1)
    if len(header_table.rows) > 1:
        values = [c.text.strip() for c in header_table.rows[1].cells]
        print(f"Values: {values}\n")
        
        # Find Sequence Number column
        seq_col = None
        for i, h in enumerate(headers):
            if 'Sequence' in h and 'Number' in h:
                seq_col = i
                print(f"✓ Found 'Sequence Number' at column {i}")
                break
        
        # Find Revision column
        rev_col = None
        for i, h in enumerate(headers):
            if 'Revision' in h:
                rev_col = i
                print(f"✓ Found 'Revision' at column {i}")
                break
        
        # Extract values
        if seq_col is not None and seq_col < len(values):
            seq = values[seq_col].strip()
            print(f"\n✓✓✓ Sequence Number = '{seq}' ✓✓✓")
        else:
            print(f"\n❌ Could not get Sequence Number value")
        
        if rev_col is not None and rev_col < len(values):
            rev = values[rev_col].strip()
            print(f"✓✓✓ Revision = '{rev}' ✓✓✓")
        else:
            print(f"❌ Could not get Revision value")


def main():
    DIR = r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\weekly"
    
    files = list(Path(DIR).glob("*.docx"))
    
    if not files:
        print("No .docx files found!")
        return
    
    # Test first file
    print("Testing FIRST file to see what's wrong...\n")
    test_one_file(files[0])


if __name__ == "__main__":
    main()
