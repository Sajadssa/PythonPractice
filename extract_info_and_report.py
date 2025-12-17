#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
استخراج اطلاعات از فایل‌های Weekly Reports و ایجاد گزارش Excel
این کد فایل‌های موجود را می‌خواند و اطلاعات آنها را استخراج می‌کند
"""

import os
from pathlib import Path
import re
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document

# سعی در import کتابخانه‌های PDF
try:
    import PyPDF2
    PDF_SUPPORT = True
except:
    PDF_SUPPORT = False
    print("⚠️ PyPDF2 نصب نیست - فقط فایل‌های Word پردازش می‌شوند")

def parse_date(date_str):
    """تبدیل تاریخ به datetime"""
    if not date_str:
        return None
    
    months = {
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12,
        'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
        'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12
    }
    
    try:
        date_str = date_str.strip()
        
        # فرمت: 15-Dec-2024
        parts = date_str.split('-')
        if len(parts) == 3:
            day = int(parts[0])
            month = months.get(parts[1].lower())
            year = int(parts[2])
            if month:
                return datetime(year, month, day)
        
        # فرمت: Dec 15, 2024
        parts = date_str.replace(',', '').split()
        if len(parts) == 3:
            month = months.get(parts[0].lower())
            day = int(parts[1])
            year = int(parts[2])
            if month:
                return datetime(year, month, day)
                
    except:
        pass
    return None

def extract_info_from_word(doc_path):
    """
    استخراج Revision و Date از فایل Word
    """
    print(f"   📄 در حال خواندن: {doc_path.name}")
    
    try:
        doc = Document(doc_path)
        
        revision = None
        date_obj = None
        date_str = None
        
        # استراتژی 1: جستجو بر اساس نام ستون‌ها
        for table_idx, table in enumerate(doc.tables):
            if revision and date_obj:
                break
            
            if len(table.rows) < 2:
                continue
            
            # خواندن هدرها
            headers = []
            for cell in table.rows[0].cells:
                header = cell.text.strip().replace('\n', ' ').replace('\r', ' ')
                headers.append(header)
            
            # پیدا کردن index ستون‌ها
            rev_idx = -1
            date_idx = -1
            
            for idx, header in enumerate(headers):
                header_lower = header.lower()
                
                if 'revision' in header_lower or header.lower() == 'rev':
                    rev_idx = idx
                
                if 'date' in header_lower:
                    date_idx = idx
            
            # استخراج داده‌ها
            for row_idx in range(1, len(table.rows)):
                data_row = table.rows[row_idx]
                
                # Revision
                if not revision and rev_idx >= 0 and rev_idx < len(data_row.cells):
                    val = data_row.cells[rev_idx].text.strip()
                    if val and len(val) >= 2 and val[0].upper() == 'G':
                        revision = val.upper()
                
                # Date
                if not date_obj and date_idx >= 0 and date_idx < len(data_row.cells):
                    val = data_row.cells[date_idx].text.strip()
                    temp_date = parse_date(val)
                    if temp_date:
                        date_obj = temp_date
                        date_str = val
                
                if revision and date_obj:
                    break
        
        # استراتژی 2: جستجو در تمام سل‌ها
        if not revision or not date_obj:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        
                        # جستجوی Revision
                        if not revision:
                            match = re.search(r'\b(G\d{2})\b', text, re.IGNORECASE)
                            if match:
                                revision = match.group(1).upper()
                        
                        # جستجوی Date
                        if not date_obj:
                            temp_date = parse_date(text)
                            if temp_date:
                                date_obj = temp_date
                                date_str = text
                        
                        if revision and date_obj:
                            break
                    if revision and date_obj:
                        break
                if revision and date_obj:
                    break
        
        print(f"      Revision: {revision if revision else '❌ یافت نشد'}")
        print(f"      Date: {date_str if date_str else '❌ یافت نشد'}")
        
        return {
            'revision': revision,
            'date': date_obj,
            'date_str': date_str
        }
        
    except Exception as e:
        print(f"      ❌ خطا: {str(e)}")
        return {
            'revision': None,
            'date': None,
            'date_str': None
        }

def extract_info_from_pdf(pdf_path):
    """استخراج Revision و Date از فایل PDF"""
    if not PDF_SUPPORT:
        return {
            'revision': None,
            'date': None,
            'date_str': None
        }
    
    print(f"   📄 در حال خواندن: {pdf_path.name}")
    
    try:
        text = ""
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            if len(reader.pages) > 0:
                text = reader.pages[0].extract_text()
        
        if not text:
            return {
                'revision': None,
                'date': None,
                'date_str': None
            }
        
        # جستجوی Revision
        revision = None
        rev_match = re.search(r'\b(G\d{2})\b', text, re.IGNORECASE)
        if rev_match:
            revision = rev_match.group(1).upper()
        
        # جستجوی Date
        date_obj = None
        date_str = None
        date_match = re.search(r'(\d{1,2}-[A-Za-z]{3,9}-\d{4})', text)
        if date_match:
            date_str = date_match.group(1)
            date_obj = parse_date(date_str)
        
        print(f"      Revision: {revision if revision else '❌ یافت نشد'}")
        print(f"      Date: {date_str if date_str else '❌ یافت نشد'}")
        
        return {
            'revision': revision,
            'date': date_obj,
            'date_str': date_str
        }
        
    except Exception as e:
        print(f"      ❌ خطا: {str(e)}")
        return {
            'revision': None,
            'date': None,
            'date_str': None
        }

def create_excel_report(files_data, output_path):
    """ایجاد گزارش Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Weekly Reports Info"
    
    # استایل‌ها
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    
    # هدر
    headers = ['ردیف', 'نام فایل', 'Revision', 'تاریخ']
    ws.append(headers)
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border
    
    # داده‌ها
    for idx, data in enumerate(files_data, start=1):
        row = [
            idx,
            data['filename'],
            data['revision'] if data['revision'] else 'N/A',
            data['date'] if data['date'] else 'N/A'
        ]
        ws.append(row)
        
        row_num = idx + 1
        for col_idx, cell in enumerate(ws[row_num], start=1):
            cell.border = border
            cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # فرمت تاریخ
            if col_idx == 4 and isinstance(cell.value, datetime):
                cell.number_format = 'DD/MM/YYYY'
    
    # تنظیم عرض ستون‌ها
    widths = [8, 50, 15, 20]
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + idx)].width = width
    
    wb.save(output_path)
    print(f"\n✅ گزارش Excel ذخیره شد: {output_path.name}")

def main():
    """تابع اصلی"""
    
    # مسیر پوشه
    FOLDER = Path(r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\weekly")
    
    # برای تست:
    # FOLDER = Path('/home/claude')
    
    if not FOLDER.exists():
        print(f"❌ پوشه پیدا نشد: {FOLDER}")
        return
    
    print("\n" + "="*80)
    print(" 📊 استخراج اطلاعات از Weekly Reports")
    print("="*80)
    print(f"📂 مسیر: {FOLDER}\n")
    
    # جمع‌آوری فایل‌ها
    pdf_files = list(FOLDER.glob('*.pdf'))
    word_files = list(FOLDER.glob('*.docx')) + list(FOLDER.glob('*.doc'))
    
    # فیلتر فایل‌های موقت
    word_files = [f for f in word_files 
                  if not f.name.startswith('~')
                  and not f.name.startswith('.')]
    
    all_files = sorted(pdf_files + word_files, key=lambda x: x.name)
    
    print(f"📁 فایل‌های پیدا شده:")
    print(f"   PDF: {len(pdf_files)}")
    print(f"   Word: {len(word_files)}")
    print(f"   جمع: {len(all_files)}\n")
    
    if not all_files:
        print("❌ فایلی پیدا نشد!")
        return
    
    print("🔍 در حال استخراج اطلاعات...")
    print("-"*80)
    
    # پردازش فایل‌ها
    files_data = []
    
    for file_path in all_files:
        # استخراج اطلاعات
        if file_path.suffix.lower() == '.pdf':
            info = extract_info_from_pdf(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            info = extract_info_from_word(file_path)
        else:
            info = {
                'revision': None,
                'date': None,
                'date_str': None
            }
        
        # ذخیره نتیجه
        files_data.append({
            'filename': file_path.name,
            'revision': info['revision'],
            'date': info['date'],
            'date_str': info['date_str']
        })
    
    print("-"*80)
    
    # ایجاد گزارش Excel
    report_path = FOLDER / f"Weekly_Reports_Info_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    create_excel_report(files_data, report_path)
    
    # خلاصه
    total = len(files_data)
    with_revision = sum(1 for d in files_data if d['revision'])
    with_date = sum(1 for d in files_data if d['date'])
    
    print(f"\n📊 خلاصه:")
    print(f"   کل فایل‌ها: {total}")
    print(f"   دارای Revision: {with_revision}")
    print(f"   دارای تاریخ: {with_date}")
    print("="*80)
    print("\n✨ تمام شد!")

if __name__ == "__main__":
    main()
