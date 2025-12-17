#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Maintenance Report File Processor v3.0 - FINAL TESTED VERSION
============================================================
Extracts: Sequence Number, Revision, Date from PDF/Word files
Renames: SJSC-GGNRSP-MADR-REMO-{seq}-{rev}.{ext}

Author: Claude AI
Date: 2025-12-17
"""

import os
import re
from pathlib import Path
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill

# Try to import PDF and Word libraries
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️  Warning: pdfplumber not installed. PDF processing will be skipped.")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  Warning: python-docx not installed. Word processing will be skipped.")


def is_temp_file(filename):
    """Check if file is a temporary file"""
    temp_patterns = ['~$', '.tmp', '.temp', '~lock']
    return any(pattern in filename for pattern in temp_patterns)


def extract_text_from_pdf(pdf_path):
    """Extract text from first page of PDF"""
    if not PDF_AVAILABLE:
        return None
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            if len(pdf.pages) > 0:
                first_page = pdf.pages[0]
                text = first_page.extract_text()
                return text if text else None
    except Exception as e:
        print(f"  ❌ Error reading PDF: {e}")
    return None


def extract_text_from_docx(docx_path):
    """Extract text from first page of Word document"""
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = DocxDocument(docx_path)
        text = ""
        
        # Extract from tables (header info is usually in tables)
        for table in doc.tables[:3]:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
            text += "\n"
        
        # Extract from paragraphs
        for para in doc.paragraphs[:20]:
            text += para.text + "\n"
        
        return text if text.strip() else None
    except Exception as e:
        print(f"  ❌ Error reading Word: {e}")
    return None


def extract_metadata(text, filename):
    """Extract Sequence Number, Revision, and Date from text"""
    if not text:
        return None
    
    metadata = {
        'sequence': None,
        'revision': None,
        'date': None
    }
    
    # Extract Sequence Number - try multiple patterns
    seq_patterns = [
        r'Sequence\s+Number\s*[:\-|]?\s*(\d+)',  # With pipes from table
        r'Number\s*[:\-|]?\s*(\d+)',
        r'Seq\s*[:\-|]?\s*(\d+)',
        r'\|\s*(\d{4})\s*\|',  # Just number in table
        r'(?:Sequence|Number).*?(\d{3,4})',  # More flexible
    ]
    
    for pattern in seq_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            metadata['sequence'] = match.group(1).zfill(4)
            break
    
    # Extract Revision - try multiple patterns
    rev_patterns = [
        r'Revision\s*[:\-|]?\s*(G\d+)',
        r'\|\s*(G\d{2,3})\s*\|',  # In table
        r'\b(G\d{2,3})\b',  # Standalone
    ]
    
    for pattern in rev_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            metadata['revision'] = match.group(1).upper()
            break
    
    # Extract Date - try multiple patterns
    date_patterns = [
        r'Date\s*[:\-|]?\s*(\d{1,2}[-/][A-Za-z]{3}[-/]\d{4})',
        r'(\d{1,2}[-/][A-Za-z]{3}[-/]\d{4})',
        r'Date.*?(\d{1,2}\s+[A-Za-z]{3}\s+\d{4})',
    ]
    
    for pattern in date_patterns:
        match = re.search(pattern, text)
        if match:
            metadata['date'] = match.group(1)
            break
    
    # Debug output
    if metadata['sequence'] or metadata['revision']:
        print(f"  ✓ Extracted - Seq: {metadata['sequence']}, Rev: {metadata['revision']}, Date: {metadata['date']}")
    else:
        print(f"  ⚠️  Could not extract data")
        # Show first 500 chars for debugging
        print(f"  📝 Text preview: {text[:500]}...")
    
    return metadata


def process_file(file_path):
    """Process a single PDF or Word file"""
    print(f"\n📖 Processing: {file_path.name}")
    
    # Extract text based on file type
    text = None
    if file_path.suffix.lower() == '.pdf':
        text = extract_text_from_pdf(file_path)
        file_type = 'PDF'
    elif file_path.suffix.lower() in ['.docx', '.doc']:
        text = extract_text_from_docx(file_path)
        file_type = 'Word'
    else:
        print(f"  ⏭️  Skipping unsupported file type")
        return None
    
    if not text:
        print(f"  ❌ Could not extract text from file")
        return None
    
    # Extract metadata
    metadata = extract_metadata(text, file_path.name)
    
    if not metadata or not metadata['sequence'] or not metadata['revision']:
        return None
    
    # Create new filename
    new_name = f"SJSC-GGNRSP-MADR-REMO-{metadata['sequence']}-{metadata['revision']}{file_path.suffix}"
    
    return {
        'original_name': file_path.name,
        'new_name': new_name,
        'date': metadata['date'] or 'N/A',
        'sequence': metadata['sequence'],
        'revision': metadata['revision'],
        'old_path': file_path,
        'file_type': file_type
    }


def process_directory(directory_path):
    """Process all PDF and Word files in directory"""
    directory = Path(directory_path)
    
    if not directory.exists():
        print(f"\n❌ Error: Directory not found!")
        print(f"   Path: {directory_path}")
        return []
    
    print(f"\n{'='*70}")
    print(f"📁 Scanning directory...")
    print(f"{'='*70}")
    
    results = []
    skipped = []
    
    # Find all files
    all_files = (
        list(directory.glob("*.pdf")) + 
        list(directory.glob("*.docx")) + 
        list(directory.glob("*.doc"))
    )
    
    print(f"\n📊 Found {len(all_files)} file(s) total")
    
    for file_path in all_files:
        # Skip temporary files
        if is_temp_file(file_path.name):
            print(f"\n⏭️  Skipping temporary: {file_path.name}")
            skipped.append(file_path.name)
            continue
        
        # Process file
        result = process_file(file_path)
        if result:
            results.append(result)
            print(f"  ✅ New name: {result['new_name']}")
    
    print(f"\n{'='*70}")
    print(f"✅ Successfully processed: {len(results)} file(s)")
    if skipped:
        print(f"⏭️  Skipped temporary files: {len(skipped)}")
    if len(all_files) - len(results) - len(skipped) > 0:
        print(f"⚠️  Failed to process: {len(all_files) - len(results) - len(skipped)} file(s)")
    print(f"{'='*70}")
    
    return results


def create_excel_report(results, output_path):
    """Create Excel report with file information"""
    print(f"\n📊 Creating Excel report...")
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "File Rename Report"
    
    # Header row
    headers = ['Original Filename', 'New Filename', 'Date', 'Sequence', 'Revision', 'Type']
    ws.append(headers)
    
    # Style header
    header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF', size=12)
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Add data rows
    for item in results:
        ws.append([
            item['original_name'],
            item['new_name'],
            item['date'],
            item['sequence'],
            item['revision'],
            item['file_type']
        ])
    
    # Column widths
    ws.column_dimensions['A'].width = 50
    ws.column_dimensions['B'].width = 45
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 15
    ws.column_dimensions['E'].width = 12
    ws.column_dimensions['F'].width = 10
    
    # Alternating row colors
    light_fill = PatternFill(start_color='F2F2F2', end_color='F2F2F2', fill_type='solid')
    for row_idx, row in enumerate(ws.iter_rows(min_row=2, max_row=ws.max_row), start=2):
        if row_idx % 2 == 0:
            for cell in row:
                cell.fill = light_fill
        # Center align certain columns
        row[2].alignment = Alignment(horizontal='center')
        row[3].alignment = Alignment(horizontal='center')
        row[4].alignment = Alignment(horizontal='center')
        row[5].alignment = Alignment(horizontal='center')
    
    # Add summary
    summary_row = ws.max_row + 2
    ws[f'A{summary_row}'] = 'Total Files Processed:'
    ws[f'B{summary_row}'] = len(results)
    ws[f'A{summary_row}'].font = Font(bold=True, size=11)
    ws[f'B{summary_row}'].font = Font(bold=True, size=11)
    
    # Save
    try:
        wb.save(output_path)
        print(f"✅ Excel report created successfully!")
        print(f"   📄 Location: {output_path}")
        return True
    except Exception as e:
        print(f"❌ Error saving Excel file: {e}")
        return False


def rename_files(results, directory_path):
    """Rename files based on extracted data"""
    directory = Path(directory_path)
    
    print(f"\n{'='*70}")
    print(f"🔄 RENAMING FILES")
    print(f"{'='*70}\n")
    
    success_count = 0
    fail_count = 0
    
    for item in results:
        old_path = item['old_path']
        new_path = directory / item['new_name']
        
        try:
            if new_path.exists():
                print(f"⚠️  Exists: {item['new_name']}")
                fail_count += 1
                continue
            
            old_path.rename(new_path)
            print(f"✅ {item['original_name']}")
            print(f"   → {item['new_name']}")
            success_count += 1
            
        except Exception as e:
            print(f"❌ Failed: {item['original_name']}")
            print(f"   Error: {e}")
            fail_count += 1
    
    print(f"\n{'='*70}")
    print(f"✅ Renamed: {success_count}/{len(results)} files")
    if fail_count > 0:
        print(f"⚠️  Failed/Skipped: {fail_count} files")
    print(f"{'='*70}")


def show_summary(results):
    """Show summary of files to be renamed"""
    print(f"\n{'='*70}")
    print(f"📋 SUMMARY OF FILES TO BE RENAMED")
    print(f"{'='*70}\n")
    
    for i, item in enumerate(results, 1):
        print(f"{i:2d}. {item['original_name']}")
        print(f"    → {item['new_name']}")
        print(f"    📅 {item['date']} | Seq: {item['sequence']} | Rev: {item['revision']}")
        print()


def main():
    """Main function"""
    
    # Configuration
    DIRECTORY_PATH = r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\monthly"
    OUTPUT_EXCEL = r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\file_rename_report.xlsx"
    
    # Header
    print("\n" + "="*70)
    print("📋 MAINTENANCE REPORT FILE PROCESSOR v3.0")
    print("="*70)
    print(f"\n📁 Directory: {DIRECTORY_PATH}")
    print(f"📊 Output Excel: {OUTPUT_EXCEL}")
    
    # Check libraries
    if not PDF_AVAILABLE:
        print("\n⚠️  WARNING: pdfplumber not installed")
        print("   Install: pip install pdfplumber")
    if not DOCX_AVAILABLE:
        print("\n⚠️  WARNING: python-docx not installed")
        print("   Install: pip install python-docx")
    
    if not PDF_AVAILABLE and not DOCX_AVAILABLE:
        print("\n❌ ERROR: No processing libraries available!")
        print("   Install: pip install pdfplumber python-docx openpyxl")
        return
    
    # Process files
    results = process_directory(DIRECTORY_PATH)
    
    if not results:
        print(f"\n{'='*70}")
        print("⚠️  NO FILES PROCESSED")
        print(f"{'='*70}")
        print("\n🔍 Troubleshooting:")
        print("1. Check directory path is correct")
        print("2. Make sure files are not open in other programs")
        print("3. Verify files contain required data in first page:")
        print("   - Sequence Number (e.g., 0003)")
        print("   - Revision (e.g., G00)")
        print("4. Install required libraries:")
        print("   pip install pdfplumber python-docx openpyxl")
        print(f"{'='*70}\n")
        return
    
    # Show summary
    show_summary(results)
    
    # Create Excel report
    if not create_excel_report(results, OUTPUT_EXCEL):
        print("\n⚠️  Excel report creation failed, but continuing...")
    
    # Ask to rename
    print(f"\n{'='*70}")
    print("❓ RENAME FILES?")
    print(f"{'='*70}")
    print(f"\n{len(results)} file(s) will be renamed.")
    print("The Excel report has been created for your review.")
    
    response = input("\nDo you want to rename the files now? (yes/no): ").strip().lower()
    
    if response in ['yes', 'y', 'بله']:
        rename_files(results, DIRECTORY_PATH)
        print("\n✅ Renaming completed!")
    else:
        print("\n✅ Files NOT renamed. Excel report is available for review.")
    
    # Final message
    print(f"\n{'='*70}")
    print("🎉 PROCESS COMPLETED!")
    print(f"{'='*70}\n")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n⚠️  Process interrupted by user (Ctrl+C)")
    except Exception as e:
        print(f"\n\n❌ Unexpected error occurred:")
        print(f"   {e}")
        print("\n📝 Please report this error with the file details.")
        import traceback
        traceback.print_exc()
