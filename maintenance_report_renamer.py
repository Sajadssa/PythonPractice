import os
from pathlib import Path
import re
from datetime import datetime
import PyPDF2
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from collections import defaultdict

# تنظیم مسیر Tesseract (در صورت نیاز)
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def parse_date_to_excel(date_str):
    """
    تبدیل تاریخ به datetime object
    ورودی: 14-Oct-2024 یا 14-October-2024
    """
    if not date_str:
        return None
    
    try:
        months = {
            'jan': 1, 'january': 1,
            'feb': 2, 'february': 2,
            'mar': 3, 'march': 3,
            'apr': 4, 'april': 4,
            'may': 5,
            'jun': 6, 'june': 6,
            'jul': 7, 'july': 7,
            'aug': 8, 'august': 8,
            'sep': 9, 'september': 9,
            'oct': 10, 'october': 10,
            'nov': 11, 'november': 11,
            'dec': 12, 'december': 12
        }
        
        # پاک کردن فضاهای اضافی
        date_str = date_str.strip()
        
        # فرمت 1: 14-Oct-2024
        parts = date_str.split('-')
        if len(parts) == 3:
            day = int(parts[0])
            month_name = parts[1].lower()
            year = int(parts[2])
            
            month = months.get(month_name)
            if month:
                return datetime(year, month, day)
        
        # فرمت 2: 14 Oct 2024
        parts = date_str.split()
        if len(parts) == 3:
            day = int(parts[0])
            month_name = parts[1].lower()
            year = int(parts[2])
            
            month = months.get(month_name)
            if month:
                return datetime(year, month, day)
                
    except Exception as e:
        print(f"   ⚠️ خطا در تبدیل تاریخ '{date_str}': {e}")
    
    return None

def extract_text_from_pdf_with_ocr(pdf_path):
    """
    استخراج متن از PDF با OCR
    """
    try:
        print(f"   🔍 تلاش برای OCR...")
        images = convert_from_path(pdf_path, first_page=1, last_page=1, dpi=300)
        
        if images:
            text = pytesseract.image_to_string(images[0], lang='eng')
            return text
    except Exception as e:
        print(f"   ⚠️ خطا در OCR: {str(e)}")
    
    return ""

def extract_info_from_pdf(pdf_path):
    """
    استخراج اطلاعات از PDF:
    - Document No (از جدول اول)
    - Date (از جدول دوم)
    - Report Title
    - Period (From ... to ...)
    """
    text = ""
    
    try:
        # خواندن مستقیم PDF
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            if len(pdf_reader.pages) > 0:
                first_page = pdf_reader.pages[0]
                text = first_page.extract_text()
        
        # اگر متن کافی نبود، از OCR استفاده کن
        if not text or len(text.strip()) < 100:
            print(f"   ⚠️ PDF اسکن شده، استفاده از OCR...")
            text = extract_text_from_pdf_with_ocr(pdf_path)
        
        if text:
            print(f"   📄 متن استخراج شده ({len(text)} کاراکتر)")
            
            # استخراج Document No
            # الگوها برای Document No در جدول اول
            patterns_docno = [
                r'Document\s*No\.?\s*:?\s*(SJSC-[A-Z]+-[A-Z]+-[A-Z]+-(\d{4})-(G\d{2}))',
                r'Document\s*Number\s*:?\s*(SJSC-[A-Z]+-[A-Z]+-[A-Z]+-(\d{4})-(G\d{2}))',
                r'Doc\s*No\.?\s*:?\s*(SJSC-[A-Z]+-[A-Z]+-[A-Z]+-(\d{4})-(G\d{2}))',
                r'(SJSC-[A-Z]+-[A-Z]+-REWK-(\d{4})-(G\d{2}))',
                r'(SJSC-[A-Z]+-[A-Z]+-[A-Z]+-(\d{4})-(G\d{2}))',
            ]
            
            doc_no = None
            doc_number = None
            rev = None
            
            for pattern in patterns_docno:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    groups = match.groups()
                    if len(groups) >= 3:
                        doc_no = groups[0]
                        doc_number = groups[1]
                        rev = groups[2]
                        print(f"   ✅ Document No پیدا شد: {doc_no}")
                        break
            
            # استخراج Date از جدول دوم
            patterns_date = [
                r'Date\s*:?\s*([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})',
                r'Date\s+([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})',
                r'Approved\s+by\s+Date\s+([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})',
            ]
            
            date_obj = None
            date_str = None
            
            for pattern in patterns_date:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    date_str = match.group(1).replace(' ', '-')
                    date_obj = parse_date_to_excel(date_str)
                    if date_obj:
                        print(f"   ✅ Date پیدا شد: {date_str} -> {date_obj.strftime('%d/%m/%Y')}")
                        break
                if date_obj:
                    break
            
            # استخراج عنوان گزارش
            report_title = None
            if 'MAINTENANCE' in text.upper() and 'WEEKLY' in text.upper():
                report_title = 'MAINTENANCE WEEKLY REPORT'
            elif 'MAINTENANCE' in text.upper() and 'MONTHLY' in text.upper():
                report_title = 'MAINTENANCE MONTHLY REPORT'
            
            # استخراج دوره گزارش (From ... to ...)
            period = None
            period_pattern = r'\(From\s+([0-9]{1,2}[-\s][A-Za-z]{3,9})\s+to\s+([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})\)'
            period_match = re.search(period_pattern, text, re.IGNORECASE)
            if period_match:
                period = f"From {period_match.group(1)} to {period_match.group(2)}"
                print(f"   ✅ Period: {period}")
            
            return {
                'doc_no': doc_no,
                'doc_number': doc_number,
                'rev': rev,
                'date': date_obj,
                'date_str': date_str,
                'report_title': report_title,
                'period': period
            }
                    
    except Exception as e:
        print(f"   ⚠️ خطا در خواندن PDF: {str(e)}")
    
    return None

def extract_info_from_word(word_path):
    """
    استخراج اطلاعات از فایل Word
    """
    try:
        print(f"   🔍 در حال خواندن Word...")
        doc = Document(word_path)
        
        # متن کامل سند
        full_text = []
        
        # متن پاراگراف‌ها
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # متن جداول - مهم برای Document No و Date
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(' | '.join(row_text))
        
        text = '\n'.join(full_text)
        print(f"   📄 متن Word استخراج شده ({len(text)} کاراکتر)")
        print(f"   🔍 نمونه متن: {text[:300].replace(chr(10), ' ')}")
        
        if not text or len(text.strip()) < 50:
            print(f"   ⚠️ متن کافی استخراج نشد!")
            return None
        
        # استخراج Document No
        patterns_docno = [
            r'Document\s*No\.?\s*:?\s*(SJSC-[A-Z]+-[A-Z]+-[A-Z]+-(\d{4})-(G\d{2}))',
            r'Document\s*Number\s*:?\s*(SJSC-[A-Z]+-[A-Z]+-[A-Z]+-(\d{4})-(G\d{2}))',
            r'(SJSC-[A-Z]+-[A-Z]+-REWK-(\d{4})-(G\d{2}))',
            r'(SJSC-[A-Z]+-[A-Z]+-[A-Z]+-(\d{4})-(G\d{2}))',
        ]
        
        doc_no = None
        doc_number = None
        rev = None
        
        for pattern in patterns_docno:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                groups = match.groups()
                if len(groups) >= 3:
                    doc_no = groups[0]
                    doc_number = groups[1]
                    rev = groups[2]
                    print(f"   ✅ Document No: {doc_no}")
                    break
        
        # استخراج Date
        patterns_date = [
            r'Date\s*:?\s*([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})',
            r'Date\s+([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})',
        ]
        
        date_obj = None
        date_str = None
        
        for pattern in patterns_date:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                date_str = match.group(1).replace(' ', '-')
                date_obj = parse_date_to_excel(date_str)
                if date_obj:
                    print(f"   ✅ Date: {date_str} -> {date_obj.strftime('%d/%m/%Y')}")
                    break
            if date_obj:
                break
        
        # عنوان گزارش
        report_title = None
        if 'MAINTENANCE' in text.upper() and 'WEEKLY' in text.upper():
            report_title = 'MAINTENANCE WEEKLY REPORT'
        elif 'MAINTENANCE' in text.upper() and 'MONTHLY' in text.upper():
            report_title = 'MAINTENANCE MONTHLY REPORT'
        
        # دوره گزارش
        period = None
        period_pattern = r'\(From\s+([0-9]{1,2}[-\s][A-Za-z]{3,9})\s+to\s+([0-9]{1,2}[-\s][A-Za-z]{3,9}[-\s][0-9]{4})\)'
        period_match = re.search(period_pattern, text, re.IGNORECASE)
        if period_match:
            period = f"From {period_match.group(1)} to {period_match.group(2)}"
        
        return {
            'doc_no': doc_no,
            'doc_number': doc_number,
            'rev': rev,
            'date': date_obj,
            'date_str': date_str,
            'report_title': report_title,
            'period': period
        }
    
    except Exception as e:
        print(f"   ⚠️ خطا در خواندن Word: {str(e)}")
    
    return None

def create_excel_report(files_data, output_path):
    """
    ایجاد گزارش اکسل
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Maintenance Reports"
    
    # استایل‌ها
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # هدرها
    headers = ['ردیف', 'نام فایل اصلی', 'نام فایل جدید', 'عنوان گزارش', 'دوره', 'Document No', 'شماره', 'REV', 'تاریخ', 'وضعیت']
    ws.append(headers)
    
    # استایل هدر
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border
    
    # داده‌ها
    for idx, data in enumerate(files_data, start=1):
        date_value = data['date'] if data['date'] else 'N/A'
        
        row = [
            idx,
            data['old_name'],
            data['new_name'],
            data['report_title'] or 'N/A',
            data['period'] or 'N/A',
            data['doc_no'] or 'N/A',
            data['doc_number'] or 'N/A',
            data['rev'] or 'N/A',
            date_value,
            data['status']
        ]
        ws.append(row)
        
        # استایل ردیف
        row_num = idx + 1
        for col_idx, cell in enumerate(ws[row_num], start=1):
            cell.border = border
            cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # فرمت تاریخ
            if col_idx == 9 and isinstance(cell.value, datetime):
                cell.number_format = 'DD/MM/YYYY'
    
    # تنظیم عرض ستون‌ها
    column_widths = [8, 40, 45, 35, 30, 40, 12, 8, 15, 20]
    for idx, width in enumerate(column_widths, start=1):
        ws.column_dimensions[chr(64 + idx)].width = width
    
    # ذخیره فایل
    wb.save(output_path)
    print(f"\n📊 فایل اکسل ایجاد شد: {output_path}")

def rename_files(folder_path):
    """
    تغییر نام فایل‌های PDF و Word
    """
    print("="*80)
    print("🔄 تغییر نام Maintenance Weekly Reports")
    print("="*80)
    print(f"📂 مسیر پوشه: {folder_path}\n")
    
    # پیدا کردن فایل‌ها
    all_files = []
    
    # PDF files
    for pdf in Path(folder_path).glob('*.pdf'):
        if not pdf.name.startswith('SJSC-GGNRSP-MADR-REWK-'):
            all_files.append(pdf)
    
    # Word files - DOCX
    for docx in Path(folder_path).glob('*.docx'):
        if not docx.name.startswith('SJSC-GGNRSP-MADR-REWK-') and not docx.name.startswith('~
    
    if not all_files:
        print("❌ هیچ فایلی پیدا نشد!")
        return
    
    print(f"📁 {len(pdf_files)} فایل PDF و {len(word_files)} فایل Word پیدا شد\n")
    print("🔍 در حال استخراج اطلاعات...")
    print("-"*80)
    
    # استخراج اطلاعات
    files_data = []
    
    for file_path in all_files:
        print(f"\n📄 پردازش: {file_path.name}")
        print(f"   📎 نوع فایل: {file_path.suffix.upper()}")
        
        if file_path.suffix.lower() == '.pdf':
            info = extract_info_from_pdf(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            info = extract_info_from_word(file_path)
        else:
            print(f"   ⚠️ نوع فایل پشتیبانی نمی‌شود!")
            continue
        
        if info and info['doc_number'] and info['rev']:
            files_data.append({
                'path': file_path,
                'old_name': file_path.name,
                'doc_no': info['doc_no'],
                'doc_number': info['doc_number'],
                'rev': info['rev'],
                'date': info['date'],
                'date_str': info['date_str'],
                'report_title': info['report_title'],
                'period': info['period'],
                'new_name': None,
                'status': 'در انتظار'
            })
        else:
            files_data.append({
                'path': file_path,
                'old_name': file_path.name,
                'doc_no': None,
                'doc_number': None,
                'rev': None,
                'date': None,
                'date_str': None,
                'report_title': None,
                'period': None,
                'new_name': 'N/A',
                'status': 'خطا - اطلاعات یافت نشد'
            })
            print(f"   ❌ نتوانستیم اطلاعات را پیدا کنیم!")
    
    print("-"*80)
    
    # شناسایی تکراری‌ها
    print("\n🔢 بررسی تکراری‌ها و تخصیص نام...")
    
    groups = defaultdict(list)
    for data in files_data:
        if data['doc_number'] and data['rev']:
            key = f"{data['doc_number']}-{data['rev']}"
            groups[key].append(data)
    
    # تخصیص نام جدید
    for key, group in groups.items():
        if len(group) == 1:
            data = group[0]
            ext = data['path'].suffix
            data['new_name'] = f"SJSC-GGNRSP-MADR-REWK-{data['doc_number']}-{data['rev']}{ext}"
        else:
            for idx, data in enumerate(group, start=1):
                ext = data['path'].suffix
                data['new_name'] = f"SJSC-GGNRSP-MADR-REWK-{data['doc_number']}-{data['rev']}_copy{idx}{ext}"
                print(f"   ⚠️ تکراری: {data['doc_number']}-{data['rev']} -> _copy{idx}")
    
    # تغییر نام فایل‌ها
    print("\n🔄 شروع تغییر نام...")
    print("-"*80)
    
    renamed_count = 0
    failed_count = 0
    
    for data in files_data:
        if data['new_name'] and data['new_name'] != 'N/A':
            old_path = data['path']
            new_name = data['new_name']
            new_path = old_path.parent / new_name
            
            if new_path.exists() and new_path != old_path:
                print(f"⚠️ فایل با این نام وجود دارد: {new_name}")
                data['status'] = 'رد شده - نام تکراری'
                failed_count += 1
                continue
            
            try:
                old_path.rename(new_path)
                renamed_count += 1
                data['status'] = '✅ موفق'
                print(f"✅ {old_path.name}")
                print(f"   ➜ {new_name}")
            except Exception as e:
                print(f"❌ خطا: {str(e)}")
                data['status'] = f'❌ خطا: {str(e)}'
                failed_count += 1
    
    print("-"*80)
    
    # ایجاد گزارش اکسل
    excel_path = Path(folder_path) / f"Maintenance_Rename_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    create_excel_report(files_data, excel_path)
    
    # خلاصه
    print(f"\n📊 نتیجه:")
    print(f"   ✅ موفق: {renamed_count}")
    print(f"   ❌ ناموفق: {failed_count}")
    print("="*80)

def main():
    """
    تابع اصلی
    """
    FOLDER_PATH = r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\weekly"
    
    if not os.path.exists(FOLDER_PATH):
        print(f"❌ خطا: پوشه پیدا نشد!")
        print(f"مسیر: {FOLDER_PATH}")
        return
    
    print("\n⚠️ هشدار: این عملیات نام فایل‌های PDF و Word را تغییر می‌دهد!")
    print("آیا مطمئن هستید؟ (y/n): ", end='')
    
    confirmation = input().lower()
    if confirmation != 'y':
        print("❌ عملیات لغو شد.")
        return
    
    rename_files(FOLDER_PATH)
    
    print("\n✨ کار تمام شد!")

if __name__ == "__main__":
    main()
):
            all_files.append(docx)
    
    # Word files - DOC
    for doc in Path(folder_path).glob('*.doc'):
        if not doc.name.startswith('SJSC-GGNRSP-MADR-REWK-') and not doc.name.startswith('~
    
    if not all_files:
        print("❌ هیچ فایلی پیدا نشد!")
        return
    
    print(f"📁 {len(pdf_files)} فایل PDF و {len(word_files)} فایل Word پیدا شد\n")
    print("🔍 در حال استخراج اطلاعات...")
    print("-"*80)
    
    # استخراج اطلاعات
    files_data = []
    
    for file_path in all_files:
        print(f"\n📄 پردازش: {file_path.name}")
        
        if file_path.suffix.lower() == '.pdf':
            info = extract_info_from_pdf(file_path)
        else:
            info = extract_info_from_word(file_path)
        
        if info and info['doc_number'] and info['rev']:
            files_data.append({
                'path': file_path,
                'old_name': file_path.name,
                'doc_no': info['doc_no'],
                'doc_number': info['doc_number'],
                'rev': info['rev'],
                'date': info['date'],
                'date_str': info['date_str'],
                'report_title': info['report_title'],
                'period': info['period'],
                'new_name': None,
                'status': 'در انتظار'
            })
        else:
            files_data.append({
                'path': file_path,
                'old_name': file_path.name,
                'doc_no': None,
                'doc_number': None,
                'rev': None,
                'date': None,
                'date_str': None,
                'report_title': None,
                'period': None,
                'new_name': 'N/A',
                'status': 'خطا - اطلاعات یافت نشد'
            })
            print(f"   ❌ نتوانستیم اطلاعات را پیدا کنیم!")
    
    print("-"*80)
    
    # شناسایی تکراری‌ها
    print("\n🔢 بررسی تکراری‌ها و تخصیص نام...")
    
    groups = defaultdict(list)
    for data in files_data:
        if data['doc_number'] and data['rev']:
            key = f"{data['doc_number']}-{data['rev']}"
            groups[key].append(data)
    
    # تخصیص نام جدید
    for key, group in groups.items():
        if len(group) == 1:
            data = group[0]
            ext = data['path'].suffix
            data['new_name'] = f"SJSC-GGNRSP-MADR-REWK-{data['doc_number']}-{data['rev']}{ext}"
        else:
            for idx, data in enumerate(group, start=1):
                ext = data['path'].suffix
                data['new_name'] = f"SJSC-GGNRSP-MADR-REWK-{data['doc_number']}-{data['rev']}_copy{idx}{ext}"
                print(f"   ⚠️ تکراری: {data['doc_number']}-{data['rev']} -> _copy{idx}")
    
    # تغییر نام فایل‌ها
    print("\n🔄 شروع تغییر نام...")
    print("-"*80)
    
    renamed_count = 0
    failed_count = 0
    
    for data in files_data:
        if data['new_name'] and data['new_name'] != 'N/A':
            old_path = data['path']
            new_name = data['new_name']
            new_path = old_path.parent / new_name
            
            if new_path.exists() and new_path != old_path:
                print(f"⚠️ فایل با این نام وجود دارد: {new_name}")
                data['status'] = 'رد شده - نام تکراری'
                failed_count += 1
                continue
            
            try:
                old_path.rename(new_path)
                renamed_count += 1
                data['status'] = '✅ موفق'
                print(f"✅ {old_path.name}")
                print(f"   ➜ {new_name}")
            except Exception as e:
                print(f"❌ خطا: {str(e)}")
                data['status'] = f'❌ خطا: {str(e)}'
                failed_count += 1
    
    print("-"*80)
    
    # ایجاد گزارش اکسل
    excel_path = Path(folder_path) / f"Maintenance_Rename_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    create_excel_report(files_data, excel_path)
    
    # خلاصه
    print(f"\n📊 نتیجه:")
    print(f"   ✅ موفق: {renamed_count}")
    print(f"   ❌ ناموفق: {failed_count}")
    print("="*80)

def main():
    """
    تابع اصلی
    """
    FOLDER_PATH = r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\weekly"
    
    if not os.path.exists(FOLDER_PATH):
        print(f"❌ خطا: پوشه پیدا نشد!")
        print(f"مسیر: {FOLDER_PATH}")
        return
    
    print("\n⚠️ هشدار: این عملیات نام فایل‌های PDF و Word را تغییر می‌دهد!")
    print("آیا مطمئن هستید؟ (y/n): ", end='')
    
    confirmation = input().lower()
    if confirmation != 'y':
        print("❌ عملیات لغو شد.")
        return
    
    rename_files(FOLDER_PATH)
    
    print("\n✨ کار تمام شد!")

if __name__ == "__main__":
    main()
):
            all_files.append(doc)
    
    pdf_files = [f for f in all_files if f.suffix.lower() == '.pdf']
    word_files = [f for f in all_files if f.suffix.lower() in ['.docx', '.doc']]
    
    if not all_files:
        print("❌ هیچ فایلی پیدا نشد!")
        return
    
    print(f"📁 {len(pdf_files)} فایل PDF و {len(word_files)} فایل Word پیدا شد\n")
    print("🔍 در حال استخراج اطلاعات...")
    print("-"*80)
    
    # استخراج اطلاعات
    files_data = []
    
    for file_path in all_files:
        print(f"\n📄 پردازش: {file_path.name}")
        
        if file_path.suffix.lower() == '.pdf':
            info = extract_info_from_pdf(file_path)
        else:
            info = extract_info_from_word(file_path)
        
        if info and info['doc_number'] and info['rev']:
            files_data.append({
                'path': file_path,
                'old_name': file_path.name,
                'doc_no': info['doc_no'],
                'doc_number': info['doc_number'],
                'rev': info['rev'],
                'date': info['date'],
                'date_str': info['date_str'],
                'report_title': info['report_title'],
                'period': info['period'],
                'new_name': None,
                'status': 'در انتظار'
            })
        else:
            files_data.append({
                'path': file_path,
                'old_name': file_path.name,
                'doc_no': None,
                'doc_number': None,
                'rev': None,
                'date': None,
                'date_str': None,
                'report_title': None,
                'period': None,
                'new_name': 'N/A',
                'status': 'خطا - اطلاعات یافت نشد'
            })
            print(f"   ❌ نتوانستیم اطلاعات را پیدا کنیم!")
    
    print("-"*80)
    
    # شناسایی تکراری‌ها
    print("\n🔢 بررسی تکراری‌ها و تخصیص نام...")
    
    groups = defaultdict(list)
    for data in files_data:
        if data['doc_number'] and data['rev']:
            key = f"{data['doc_number']}-{data['rev']}"
            groups[key].append(data)
    
    # تخصیص نام جدید
    for key, group in groups.items():
        if len(group) == 1:
            data = group[0]
            ext = data['path'].suffix
            data['new_name'] = f"SJSC-GGNRSP-MADR-REWK-{data['doc_number']}-{data['rev']}{ext}"
        else:
            for idx, data in enumerate(group, start=1):
                ext = data['path'].suffix
                data['new_name'] = f"SJSC-GGNRSP-MADR-REWK-{data['doc_number']}-{data['rev']}_copy{idx}{ext}"
                print(f"   ⚠️ تکراری: {data['doc_number']}-{data['rev']} -> _copy{idx}")
    
    # تغییر نام فایل‌ها
    print("\n🔄 شروع تغییر نام...")
    print("-"*80)
    
    renamed_count = 0
    failed_count = 0
    
    for data in files_data:
        if data['new_name'] and data['new_name'] != 'N/A':
            old_path = data['path']
            new_name = data['new_name']
            new_path = old_path.parent / new_name
            
            if new_path.exists() and new_path != old_path:
                print(f"⚠️ فایل با این نام وجود دارد: {new_name}")
                data['status'] = 'رد شده - نام تکراری'
                failed_count += 1
                continue
            
            try:
                old_path.rename(new_path)
                renamed_count += 1
                data['status'] = '✅ موفق'
                print(f"✅ {old_path.name}")
                print(f"   ➜ {new_name}")
            except Exception as e:
                print(f"❌ خطا: {str(e)}")
                data['status'] = f'❌ خطا: {str(e)}'
                failed_count += 1
    
    print("-"*80)
    
    # ایجاد گزارش اکسل
    excel_path = Path(folder_path) / f"Maintenance_Rename_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    create_excel_report(files_data, excel_path)
    
    # خلاصه
    print(f"\n📊 نتیجه:")
    print(f"   ✅ موفق: {renamed_count}")
    print(f"   ❌ ناموفق: {failed_count}")
    print("="*80)

def main():
    """
    تابع اصلی
    """
    FOLDER_PATH = r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\weekly"
    
    if not os.path.exists(FOLDER_PATH):
        print(f"❌ خطا: پوشه پیدا نشد!")
        print(f"مسیر: {FOLDER_PATH}")
        return
    
    print("\n⚠️ هشدار: این عملیات نام فایل‌های PDF و Word را تغییر می‌دهد!")
    print("آیا مطمئن هستید؟ (y/n): ", end='')
    
    confirmation = input().lower()
    if confirmation != 'y':
        print("❌ عملیات لغو شد.")
        return
    
    rename_files(FOLDER_PATH)
    
    print("\n✨ کار تمام شد!")

if __name__ == "__main__":
    main()
