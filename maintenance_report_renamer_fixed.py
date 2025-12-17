import os
from pathlib import Path
import re
from datetime import datetime
import PyPDF2
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from collections import defaultdict

# تنظیم مسیر Tesseract (در صورت نیاز)
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def parse_date_to_excel(date_str):
    """
    تبدیل تاریخ به datetime object
    ورودی: 14-Oct-2024
    """
    if not date_str:
        return None
    
    try:
        months = {
            'jan': 1, 'january': 1,
            'feb': 2, 'february': 2,
            'mar': 3, 'march': 3,
            'apr': 4, 'april': 4,
            'may': 5,
            'jun': 6, 'june': 6,
            'jul': 7, 'july': 7,
            'aug': 8, 'august': 8,
            'sep': 9, 'september': 9,
            'oct': 10, 'october': 10,
            'nov': 11, 'november': 11,
            'dec': 12, 'december': 12
        }
        
        # حذف فضاهای اضافی
        date_str = date_str.strip()
        
        # فرمت: 14-Oct-2024 یا 14-October-2024
        parts = date_str.split('-')
        if len(parts) == 3:
            day = int(parts[0])
            month_name = parts[1].lower()
            year = int(parts[2])
            
            month = months.get(month_name)
            if month:
                return datetime(year, month, day)
    except:
        pass
    
    return None

def extract_text_from_pdf_with_ocr(pdf_path):
    """
    استخراج متن از PDF با OCR
    """
    try:
        print(f"   🔍 استفاده از OCR برای PDF اسکن شده...")
        images = convert_from_path(pdf_path, first_page=1, last_page=1, dpi=300)
        
        if images:
            text = pytesseract.image_to_string(images[0], lang='eng')
            return text
    except Exception as e:
        print(f"   ⚠️ خطا در OCR: {str(e)}")
    
    return ""

def extract_document_number_from_table(text):
    """
    استخراج شماره سند از جدول هدر
    جدول شامل: Project Code | Scope | Discipline | Document Type | Sequence Number | Revision
    مثال: SJSC | GGNRSP | PDME | REWK | 0047 | G00
    """
    # الگو برای یافتن جدول و استخراج اطلاعات
    # جستجوی الگوی: SJSC ... GGNRSP ... (discipline) ... REWK ... (number) ... G##
    
    patterns = [
        # الگو 1: جدول کامل
        r'SJSC[|\s]+GGNRSP[|\s]+[A-Z]+[|\s]+REWK[|\s]+(\d{4})[|\s]+(G\d{2})',
        # الگو 2: فقط بخش‌های مهم
        r'REWK[|\s]+(\d{4})[|\s]+(G\d{2})',
        # الگو 3: Document Type و Sequence
        r'Document\s+Type.*?REWK.*?(\d{4}).*?(G\d{2})',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            sequence_number = match.group(1)
            revision = match.group(2)
            return sequence_number, revision
    
    return None, None

def extract_date_from_table(text):
    """
    استخراج تاریخ از جدول دوم (Rev table)
    جستجوی Date در جدول با ستون‌های: Rev | Purpose for Review | ... | Date
    """
    patterns = [
        # الگو 1: Date با عدد
        r'Date[|\s:]*([0-9]{1,2}-[A-Za-z]{3,9}-[0-9]{4})',
        # الگو 2: در کنار G00 یا IFI
        r'G00.*?([0-9]{1,2}-[A-Za-z]{3,9}-[0-9]{4})',
        r'IFI.*?([0-9]{1,2}-[A-Za-z]{3,9}-[0-9]{4})',
        # الگو 3: فقط تاریخ
        r'([0-9]{1,2}-[A-Za-z]{3,9}-[0-9]{4})',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            date_str = match.group(1)
            date_obj = parse_date_to_excel(date_str)
            if date_obj:
                return date_obj, date_str
    
    return None, None

def extract_info_from_pdf(pdf_path):
    """
    استخراج اطلاعات از PDF
    """
    text = ""
    
    try:
        # خواندن متن مستقیم
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            if len(pdf_reader.pages) > 0:
                first_page = pdf_reader.pages[0]
                text = first_page.extract_text()
        
        # اگر متن خالی بود، از OCR استفاده کن
        if not text or len(text.strip()) < 50:
            text = extract_text_from_pdf_with_ocr(pdf_path)
        
        if text:
            print(f"   📄 متن استخراج شده ({len(text)} کاراکتر)")
            
            # استخراج شماره سند
            sequence_number, revision = extract_document_number_from_table(text)
            
            # استخراج تاریخ
            date_obj, date_str = extract_date_from_table(text)
            
            if sequence_number and revision:
                print(f"   ✅ Sequence: {sequence_number}, Rev: {revision}")
            if date_obj:
                print(f"   ✅ Date: {date_str}")
            
            return {
                'sequence_number': sequence_number,
                'revision': revision,
                'date': date_obj,
                'date_str': date_str
            }
                    
    except Exception as e:
        print(f"   ⚠️ خطا در خواندن PDF: {str(e)}")
    
    return None

def extract_info_from_word(word_path):
    """
    استخراج اطلاعات از فایل Word
    """
    try:
        doc = Document(word_path)
        
        # استخراج متن از جداول
        text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
                text += "\n"
        
        # استخراج متن از پاراگراف‌ها
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        print(f"   📄 متن Word استخراج شده ({len(text)} کاراکتر)")
        
        # استخراج شماره سند
        sequence_number, revision = extract_document_number_from_table(text)
        
        # استخراج تاریخ
        date_obj, date_str = extract_date_from_table(text)
        
        if sequence_number and revision:
            print(f"   ✅ Sequence: {sequence_number}, Rev: {revision}")
        if date_obj:
            print(f"   ✅ Date: {date_str}")
        
        return {
            'sequence_number': sequence_number,
            'revision': revision,
            'date': date_obj,
            'date_str': date_str
        }
    
    except Exception as e:
        print(f"   ⚠️ خطا در خواندن Word: {str(e)}")
    
    return None

def create_excel_report(files_data, output_path):
    """
    ایجاد گزارش اکسل
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Maintenance Weekly Reports"
    
    # استایل‌ها
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # هدرها
    headers = ['ردیف', 'نام فایل اصلی', 'نام فایل جدید', 'Sequence Number', 'Revision', 'تاریخ', 'وضعیت']
    ws.append(headers)
    
    # استایل هدر
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border
    
    # داده‌ها
    for idx, data in enumerate(files_data, start=1):
        date_value = data['date'] if data['date'] else 'N/A'
        
        row = [
            idx,
            data['old_name'],
            data['new_name'],
            data['sequence_number'] or 'N/A',
            data['revision'] or 'N/A',
            date_value,
            data['status']
        ]
        ws.append(row)
        
        # استایل ردیف
        row_num = idx + 1
        for col_idx, cell in enumerate(ws[row_num], start=1):
            cell.border = border
            cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # فرمت تاریخ
            if col_idx == 6 and isinstance(cell.value, datetime):
                cell.number_format = 'DD/MM/YYYY'
    
    # تنظیم عرض ستون‌ها
    column_widths = [8, 40, 45, 18, 12, 15, 20]
    for idx, width in enumerate(column_widths, start=1):
        ws.column_dimensions[chr(64 + idx)].width = width
    
    # ذخیره فایل
    wb.save(output_path)
    print(f"\n📊 فایل اکسل ایجاد شد: {output_path}")

def rename_files(folder_path):
    """
    تغییر نام فایل‌های PDF و Word
    """
    print("="*80)
    print("🔄 تغییر نام Maintenance Weekly Reports")
    print("="*80)
    print(f"📂 مسیر پوشه: {folder_path}\n")
    
    # پیدا کردن فایل‌ها
    pdf_files = list(Path(folder_path).glob('*.pdf'))
    word_files = list(Path(folder_path).glob('*.docx')) + list(Path(folder_path).glob('*.doc'))
    
    # حذف فایل‌هایی که قبلاً تغییر نام داده شده‌اند و فایل‌های موقت
    pdf_files = [f for f in pdf_files if not f.name.startswith('SJSC-GGNRSP-MADR-REWK-')]
    word_files = [f for f in word_files if not f.name.startswith('SJSC-GGNRSP-MADR-REWK-') and not f.name.startswith('~')]
    
    all_files = pdf_files + word_files
    
    if not all_files:
        print("❌ هیچ فایلی پیدا نشد!")
        return
    
    print(f"📁 {len(pdf_files)} فایل PDF و {len(word_files)} فایل Word پیدا شد\n")
    print("🔍 در حال استخراج اطلاعات...")
    print("-"*80)
    
    # استخراج اطلاعات
    files_data = []
    
    for file_path in all_files:
        print(f"\n📄 پردازش: {file_path.name}")
        
        if file_path.suffix.lower() == '.pdf':
            info = extract_info_from_pdf(file_path)
        else:
            info = extract_info_from_word(file_path)
        
        if info and info['sequence_number'] and info['revision']:
            files_data.append({
                'path': file_path,
                'old_name': file_path.name,
                'sequence_number': info['sequence_number'],
                'revision': info['revision'],
                'date': info['date'],
                'date_str': info['date_str'],
                'new_name': None,  # بعداً پر می‌شود
                'status': 'در انتظار'
            })
        else:
            files_data.append({
                'path': file_path,
                'old_name': file_path.name,
                'sequence_number': None,
                'revision': None,
                'date': None,
                'date_str': None,
                'new_name': 'N/A',
                'status': 'خطا - اطلاعات یافت نشد'
            })
            print(f"   ❌ نتوانستیم اطلاعات را پیدا کنیم!")
    
    print("-"*80)
    
    # شناسایی تکراری‌ها و تخصیص نام
    print("\n🔢 بررسی تکراری‌ها و تخصیص نام...")
    
    # گروه‌بندی بر اساس شماره + REV
    groups = defaultdict(list)
    for data in files_data:
        if data['sequence_number'] and data['revision']:
            key = f"{data['sequence_number']}-{data['revision']}"
            groups[key].append(data)
    
    # تخصیص نام جدید
    for key, group in groups.items():
        if len(group) == 1:
            # فایل یکتا
            data = group[0]
            ext = data['path'].suffix
            data['new_name'] = f"SJSC-GGNRSP-MADR-REWK-{data['sequence_number']}-{data['revision']}{ext}"
        else:
            # فایل‌های تکراری
            for idx, data in enumerate(group, start=1):
                ext = data['path'].suffix
                data['new_name'] = f"SJSC-GGNRSP-MADR-REWK-{data['sequence_number']}-{data['revision']}_copy{idx}{ext}"
                print(f"   ⚠️ تکراری: {data['sequence_number']}-{data['revision']} -> _copy{idx}")
    
    # تغییر نام فایل‌ها
    print("\n🔄 شروع تغییر نام...")
    print("-"*80)
    
    renamed_count = 0
    failed_count = 0
    
    for data in files_data:
        if data['new_name'] and data['new_name'] != 'N/A':
            old_path = data['path']
            new_name = data['new_name']
            new_path = old_path.parent / new_name
            
            if new_path.exists() and new_path != old_path:
                print(f"⚠️ فایل با این نام وجود دارد: {new_name}")
                data['status'] = 'رد شده - نام تکراری'
                failed_count += 1
                continue
            
            try:
                old_path.rename(new_path)
                renamed_count += 1
                data['status'] = '✅ موفق'
                print(f"✅ {old_path.name}")
                print(f"   ➜ {new_name}")
            except Exception as e:
                print(f"❌ خطا: {str(e)}")
                data['status'] = f'❌ خطا: {str(e)}'
                failed_count += 1
    
    print("-"*80)
    
    # ایجاد گزارش اکسل
    excel_path = Path(folder_path) / f"Weekly_Rename_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    create_excel_report(files_data, excel_path)
    
    # خلاصه
    print(f"\n📊 نتیجه:")
    print(f"   ✅ تعداد فایل‌های تغییر نام داده شده: {renamed_count}")
    print(f"   ❌ تعداد فایل‌های ناموفق: {failed_count}")
    print("="*80)

def main():
    """
    تابع اصلی
    """
    FOLDER_PATH = r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\weekly"
    
    if not os.path.exists(FOLDER_PATH):
        print(f"❌ خطا: پوشه پیدا نشد!")
        print(f"مسیر: {FOLDER_PATH}")
        return
    
    print("\n⚠️ هشدار: این عملیات نام فایل‌های PDF و Word را تغییر می‌دهد!")
    print("⚠️ برای PDF های اسکن شده از OCR استفاده می‌شود")
    print("\nآیا مطمئن هستید؟ (y/n): ", end='')
    
    confirmation = input().lower()
    if confirmation != 'y':
        print("❌ عملیات لغو شد.")
        return
    
    rename_files(FOLDER_PATH)
    
    print("\n✨ کار تمام شد!")

if __name__ == "__main__":
    main()
