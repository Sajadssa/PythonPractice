import os
import re
from pathlib import Path
import pandas as pd
from datetime import datetime

# کتابخانه های مورد نیاز برای خواندن PDF و Word
try:
    import PyPDF2
    from pdf2image import convert_from_path
    import pytesseract
    from PIL import Image
except ImportError:
    print("در حال نصب کتابخانه‌های PDF...")

try:
    from docx import Document
except ImportError:
    print("در حال نصب کتابخانه‌های Word...")

def extract_info_from_pdf(pdf_path):
    """استخراج اطلاعات از فایل PDF"""
    try:
        # روش 1: خواندن متن مستقیم از PDF
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            if len(pdf_reader.pages) > 0:
                first_page = pdf_reader.pages[0]
                text = first_page.extract_text()
                
                # جستجوی Sequence Number
                seq_match = re.search(r'Sequence\s*Number\s*[:\s]*(\d+)', text, re.IGNORECASE)
                sequence_number = seq_match.group(1) if seq_match else None
                
                # جستجوی Revision
                rev_match = re.search(r'Revision\s*[:\s]*(G\d+)', text, re.IGNORECASE)
                revision = rev_match.group(1) if rev_match else None
                
                # جستجوی تاریخ (فرمت‌های مختلف)
                date_patterns = [
                    r'(\d{2}[-/]\w{3}[-/]\d{4})',  # 05-Feb-2024
                    r'(\d{2}[-/]\d{2}[-/]\d{4})',   # 05/02/2024
                    r'(\w{3}\s+\d{1,2},?\s+\d{4})'  # Feb 05, 2024
                ]
                
                date_str = None
                for pattern in date_patterns:
                    date_match = re.search(pattern, text)
                    if date_match:
                        date_str = date_match.group(1)
                        break
                
                return sequence_number, revision, date_str
        
        return None, None, None
        
    except Exception as e:
        print(f"خطا در خواندن PDF {pdf_path}: {str(e)}")
        return None, None, None

def extract_info_from_word(word_path):
    """استخراج اطلاعات از فایل Word"""
    try:
        doc = Document(word_path)
        full_text = ""
        
        # خواندن تمام پاراگراف‌ها
        for para in doc.paragraphs:
            full_text += para.text + "\n"
        
        # خواندن جداول (معمولاً اطلاعات در جدول صفحه اول است)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += cell.text + "\n"
        
        # جستجوی Sequence Number
        seq_match = re.search(r'Sequence\s*Number\s*[:\s]*(\d+)', full_text, re.IGNORECASE)
        sequence_number = seq_match.group(1) if seq_match else None
        
        # جستجوی Revision
        rev_match = re.search(r'Revision\s*[:\s]*(G\d+)', full_text, re.IGNORECASE)
        revision = rev_match.group(1) if rev_match else None
        
        # جستجوی تاریخ
        date_patterns = [
            r'(\d{2}[-/]\w{3}[-/]\d{4})',
            r'(\d{2}[-/]\d{2}[-/]\d{4})',
            r'(\w{3}\s+\d{1,2},?\s+\d{4})'
        ]
        
        date_str = None
        for pattern in date_patterns:
            date_match = re.search(pattern, full_text)
            if date_match:
                date_str = date_match.group(1)
                break
        
        return sequence_number, revision, date_str
        
    except Exception as e:
        print(f"خطا در خواندن Word {word_path}: {str(e)}")
        return None, None, None

def format_sequence_number(seq_num):
    """فرمت کردن شماره sequence به 4 رقمی"""
    if seq_num:
        return seq_num.zfill(4)
    return None

def generate_new_filename(sequence_number, revision):
    """تولید نام جدید فایل"""
    if sequence_number and revision:
        formatted_seq = format_sequence_number(sequence_number)
        return f"SJSC-GGNRSP-MADR-REMO-{formatted_seq}-{revision}"
    return None

def process_files(directory_path):
    """پردازش تمام فایل‌های PDF و Word در دایرکتوری"""
    
    # بررسی وجود دایرکتوری
    if not os.path.exists(directory_path):
        print(f"خطا: دایرکتوری {directory_path} وجود ندارد!")
        return
    
    results = []
    processed_count = 0
    error_count = 0
    
    # لیست فایل‌های PDF و Word
    files = list(Path(directory_path).glob("*.pdf")) + list(Path(directory_path).glob("*.docx"))
    
    print(f"\n{'='*80}")
    print(f"تعداد کل فایل‌های یافت شده: {len(files)}")
    print(f"{'='*80}\n")
    
    for file_path in files:
        print(f"در حال پردازش: {file_path.name}")
        
        # استخراج اطلاعات بر اساس نوع فایل
        if file_path.suffix.lower() == '.pdf':
            sequence_number, revision, date_str = extract_info_from_pdf(str(file_path))
        elif file_path.suffix.lower() == '.docx':
            sequence_number, revision, date_str = extract_info_from_word(str(file_path))
        else:
            continue
        
        # بررسی موفقیت استخراج
        if sequence_number and revision:
            new_filename = generate_new_filename(sequence_number, revision)
            new_filepath = file_path.parent / f"{new_filename}{file_path.suffix}"
            
            # تغییر نام فایل
            try:
                # اگر فایل با نام جدید وجود داشت، آن را حذف کن
                if new_filepath.exists():
                    print(f"  ⚠️  فایل {new_filepath.name} قبلاً وجود دارد و جایگزین می‌شود")
                
                file_path.rename(new_filepath)
                print(f"  ✓ نام قدیم: {file_path.name}")
                print(f"  ✓ نام جدید: {new_filename}{file_path.suffix}")
                print(f"  ✓ تاریخ: {date_str if date_str else 'یافت نشد'}")
                
                results.append({
                    'نام قدیم': file_path.name,
                    'نام جدید': f"{new_filename}{file_path.suffix}",
                    'Sequence Number': sequence_number,
                    'Revision': revision,
                    'تاریخ': date_str if date_str else 'N/A'
                })
                processed_count += 1
                
            except Exception as e:
                print(f"  ✗ خطا در تغییر نام: {str(e)}")
                error_count += 1
        else:
            print(f"  ✗ اطلاعات کامل یافت نشد:")
            print(f"    - Sequence Number: {sequence_number if sequence_number else 'یافت نشد'}")
            print(f"    - Revision: {revision if revision else 'یافت نشد'}")
            error_count += 1
        
        print()
    
    # ذخیره نتایج در Excel
    if results:
        df = pd.DataFrame(results)
        excel_path = Path(directory_path) / 'تغییرات_نام_فایل‌ها.xlsx'
        df.to_excel(excel_path, index=False, engine='openpyxl')
        print(f"\n{'='*80}")
        print(f"✓ فایل Excel با موفقیت ایجاد شد: {excel_path}")
        print(f"✓ تعداد فایل‌های پردازش شده: {processed_count}")
        if error_count > 0:
            print(f"⚠️  تعداد فایل‌های با خطا: {error_count}")
        print(f"{'='*80}\n")
        
        return excel_path
    else:
        print("\n⚠️  هیچ فایلی با موفقیت پردازش نشد!")
        return None

if __name__ == "__main__":
    # مسیر دایرکتوری را وارد کنید
    directory = r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\monthly"
    
    print("\n" + "="*80)
    print("برنامه تغییر نام و استخراج اطلاعات فایل‌های گزارش نگهداری")
    print("="*80)
    
    excel_file = process_files(directory)
    
    if excel_file:
        print(f"\n✓ عملیات با موفقیت انجام شد!")
        print(f"✓ فایل Excel در مسیر زیر ذخیره شد:")
        print(f"  {excel_file}")
    else:
        print("\n✗ عملیات با خطا مواجه شد!")
