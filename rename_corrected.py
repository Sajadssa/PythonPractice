#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
برنامه تصحیح نام فایل‌های Word گزارشات هفتگی
این برنامه Sequence Number را از داخل فایل می‌خواند و نام‌های اشتباه را تصحیح می‌کند
"""

import os
import re
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill

def extract_sequence_number(doc):
    """استخراج Sequence Number از داخل سند"""
    try:
        # جستجو در تمام جداول
        for table in doc.tables:
            # ابتدا بررسی می‌کنیم آیا این جدول ستون Sequence Number دارد
            header_row = None
            seq_col_idx = None
            
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                
                # پیدا کردن ستون Sequence Number
                for col_idx, cell_text in enumerate(cells):
                    if 'Sequence' in cell_text and 'Number' in cell_text:
                        header_row = row_idx
                        seq_col_idx = col_idx
                        break
                
                if seq_col_idx is not None:
                    break
            
            # اگر ستون Sequence Number پیدا شد، مقدار آن را بخوان
            if seq_col_idx is not None and header_row is not None:
                # مقدار در ردیف بعدی همان ستون
                if header_row + 1 < len(table.rows):
                    next_row = table.rows[header_row + 1]
                    if seq_col_idx < len(next_row.cells):
                        seq = next_row.cells[seq_col_idx].text.strip()
                        # حذف فضاهای خالی و کاراکترهای اضافی
                        seq = re.sub(r'\D', '', seq)  # فقط اعداد
                        if seq and seq.isdigit():
                            return seq.zfill(4)
        
        # اگر پیدا نشد، جستجوی عمومی‌تر
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    # اگر سلول فقط یک عدد 4 رقمی است (و سال نیست)
                    if text.isdigit() and 1 <= len(text) <= 4:
                        num = int(text)
                        if num > 0 and num < 10000 and text not in ['2024', '2025', '2026']:
                            return text.zfill(4)
        
    except Exception as e:
        print(f"    ⚠️ خطا در استخراج: {e}")
    
    return None

def extract_revision(doc):
    """استخراج Revision از داخل سند"""
    try:
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                
                for col_idx, cell_text in enumerate(cells):
                    if 'Revision' in cell_text:
                        # مقدار در ردیف بعدی
                        if row_idx + 1 < len(table.rows):
                            next_row = table.rows[row_idx + 1]
                            if col_idx < len(next_row.cells):
                                rev = next_row.cells[col_idx].text.strip()
                                # الگوی G + دو رقم
                                match = re.search(r'G\d{2}', rev, re.IGNORECASE)
                                if match:
                                    return match.group(0).upper()
        
        # جستجوی کلی
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    match = re.search(r'\b(G\d{2})\b', text, re.IGNORECASE)
                    if match:
                        return match.group(1).upper()
                        
    except Exception as e:
        print(f"    ⚠️ خطا در استخراج Revision: {e}")
    
    return "G00"

def extract_date(doc):
    """استخراج تاریخ از داخل سند"""
    try:
        # جستجو در جدول اول (معمولاً جدول اطلاعات تایید)
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                
                for col_idx, cell_text in enumerate(cells):
                    if cell_text == 'Date':
                        # مقدار در ردیف بعدی
                        if row_idx + 1 < len(table.rows):
                            next_row = table.rows[row_idx + 1]
                            if col_idx < len(next_row.cells):
                                date = next_row.cells[col_idx].text.strip()
                                if date and date != 'Date':
                                    return date
        
        # جستجوی الگوی تاریخ در تمام جداول
        date_pattern = r'\d{1,2}[-/\.]\w{3}[-/\.]\d{4}'
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    match = re.search(date_pattern, cell.text)
                    if match:
                        return match.group(0)
                        
    except Exception as e:
        print(f"    ⚠️ خطا در استخراج تاریخ: {e}")
    
    return "N/A"

def get_current_sequence_from_filename(filename):
    """استخراج Sequence Number فعلی از نام فایل (اگر وجود دارد)"""
    pattern = r'SJSC-GGNRSP-(?:EPWC|PDOP)-REWK-(\d{4})-'
    match = re.search(pattern, filename, re.IGNORECASE)
    if match:
        return match.group(1)
    return None

def process_file(file_path, debug=False):
    """پردازش یک فایل و استخراج اطلاعات"""
    filename = os.path.basename(file_path)
    print(f"\n{'='*70}")
    print(f"📄 {filename}")
    
    try:
        # خواندن سند
        doc = Document(file_path)
        
        if debug:
            print(f"\n  🔍 ساختار سند:")
            print(f"  تعداد جداول: {len(doc.tables)}")
            for idx, table in enumerate(doc.tables):
                print(f"\n  جدول {idx + 1}: {len(table.rows)} ردیف × {len(table.columns)} ستون")
                for i, row in enumerate(table.rows[:3]):
                    cells = [cell.text.strip()[:40] for cell in row.cells]
                    print(f"    ردیف {i+1}: {cells}")
        
        # استخراج اطلاعات از داخل سند
        correct_sequence = extract_sequence_number(doc)
        revision = extract_revision(doc)
        date = extract_date(doc)
        
        if not correct_sequence:
            print(f"  ❌ Sequence Number در داخل سند یافت نشد!")
            print(f"  💡 از debug mode استفاده کنید: python script.py --debug")
            return None
        
        # Sequence Number فعلی در نام فایل
        current_sequence = get_current_sequence_from_filename(filename)
        
        # ساخت نام صحیح
        correct_name = f"SJSC-GGNRSP-EPWC-REWK-{correct_sequence}-{revision}.docx"
        
        # بررسی آیا نام فعلی درست است یا خیر
        needs_rename = (filename != correct_name)
        
        print(f"\n  📊 اطلاعات استخراج شده:")
        print(f"  ├─ Sequence Number (صحیح): {correct_sequence}")
        if current_sequence:
            print(f"  ├─ Sequence در نام فایل: {current_sequence}")
            if current_sequence != correct_sequence:
                print(f"  ├─ ⚠️  اختلاف: {current_sequence} → {correct_sequence}")
        print(f"  ├─ Revision: {revision}")
        print(f"  └─ تاریخ: {date}")
        
        print(f"\n  📝 نام صحیح: {correct_name}")
        
        if needs_rename:
            print(f"  ✅ نیاز به تصحیح دارد")
        else:
            print(f"  ✓ نام فایل صحیح است")
        
        return {
            'old_name': filename,
            'new_name': correct_name,
            'current_sequence': current_sequence or 'N/A',
            'correct_sequence': correct_sequence,
            'revision': revision,
            'date': date,
            'needs_rename': needs_rename,
            'old_path': file_path
        }
        
    except Exception as e:
        print(f"  ❌ خطا در پردازش: {e}")
        import traceback
        traceback.print_exc()
        return None

def create_excel_report(results, output_path):
    """ایجاد گزارش Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "گزارش تصحیح نام‌ها"
    
    # هدرها
    headers = ['ردیف', 'نام فعلی', 'نام صحیح', 'Seq فعلی', 'Seq صحیح', 'Revision', 'تاریخ', 'وضعیت']
    ws.append(headers)
    
    # فرمت هدر
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # داده‌ها
    for idx, r in enumerate(results, 1):
        status = "🔴 نیاز به تصحیح" if r['needs_rename'] else "✅ صحیح"
        
        row_data = [
            idx,
            r['old_name'],
            r['new_name'],
            r['current_sequence'],
            r['correct_sequence'],
            r['revision'],
            r['date'],
            status
        ]
        ws.append(row_data)
        
        # رنگ‌بندی
        row_num = idx + 1
        if r['needs_rename']:
            # رنگ قرمز برای فایل‌هایی که نیاز به تصحیح دارند
            for col in range(1, 9):
                ws.cell(row_num, col).fill = PatternFill(start_color="FFE6E6", end_color="FFE6E6", fill_type="solid")
        
        for cell in ws[row_num]:
            cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # عرض ستون‌ها
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 45
    ws.column_dimensions['C'].width = 40
    ws.column_dimensions['D'].width = 12
    ws.column_dimensions['E'].width = 12
    ws.column_dimensions['F'].width = 10
    ws.column_dimensions['G'].width = 15
    ws.column_dimensions['H'].width = 18
    
    wb.save(output_path)
    print(f"\n✅ گزارش Excel ذخیره شد: {output_path}")

def rename_files(results):
    """تصحیح نام فایل‌ها"""
    files_to_rename = [r for r in results if r['needs_rename']]
    
    if not files_to_rename:
        print("\n✅ همه فایل‌ها نام صحیح دارند!")
        return
    
    print(f"\n{'='*70}")
    print(f"🔄 تصحیح {len(files_to_rename)} فایل...")
    print(f"{'='*70}\n")
    
    success = 0
    failed = 0
    
    for r in files_to_rename:
        old_path = r['old_path']
        new_path = os.path.join(os.path.dirname(old_path), r['new_name'])
        
        try:
            if os.path.exists(new_path):
                print(f"⚠️  فایل موجود است: {r['new_name']}")
                # اضافه کردن شماره
                base, ext = os.path.splitext(new_path)
                counter = 1
                while os.path.exists(f"{base}_copy{counter}{ext}"):
                    counter += 1
                new_path = f"{base}_copy{counter}{ext}"
                r['new_name'] = os.path.basename(new_path)
            
            os.rename(old_path, new_path)
            print(f"✅ {r['old_name']}")
            print(f"   → {r['new_name']}\n")
            success += 1
            
        except Exception as e:
            print(f"❌ خطا: {r['old_name']}")
            print(f"   {e}\n")
            failed += 1
    
    print(f"{'='*70}")
    print(f"نتیجه: ✅ {success} موفق، ❌ {failed} ناموفق")
    print(f"{'='*70}")

def main():
    """تابع اصلی"""
    import sys
    
    debug_mode = '--debug' in sys.argv
    
    print("="*70)
    print("🔧 برنامه تصحیح نام فایل‌های گزارشات هفتگی")
    print("="*70)
    
    # دریافت مسیر پوشه
    if len(sys.argv) > 1 and not sys.argv[1].startswith('--'):
        folder = sys.argv[1]
    else:
        folder = input("\n📁 مسیر پوشه فایل‌ها را وارد کنید\n(یا Enter برای پوشه فعلی): ").strip()
        if not folder:
            folder = os.getcwd()
    
    if not os.path.exists(folder):
        print(f"\n❌ پوشه یافت نشد: {folder}")
        return
    
    # پیدا کردن فایل‌ها
    files = [f for f in os.listdir(folder) 
             if f.endswith('.docx') and not f.startswith('~$')]
    
    if not files:
        print("\n❌ هیچ فایل Word یافت نشد!")
        return
    
    print(f"\n📊 تعداد فایل‌های یافت شده: {len(files)}")
    
    if debug_mode:
        print("\n🔍 Debug Mode فعال است")
    
    # پردازش فایل‌ها
    results = []
    for f in sorted(files):
        result = process_file(os.path.join(folder, f), debug=debug_mode)
        if result:
            results.append(result)
    
    if not results:
        print("\n❌ هیچ فایلی با موفقیت پردازش نشد!")
        print("💡 از debug mode استفاده کنید:")
        print("   python script.py --debug")
        return
    
    # آمار
    needs_rename = sum(1 for r in results if r['needs_rename'])
    print(f"\n{'='*70}")
    print(f"📈 خلاصه:")
    print(f"  ├─ کل فایل‌ها: {len(results)}")
    print(f"  ├─ نیاز به تصحیح: {needs_rename}")
    print(f"  └─ صحیح: {len(results) - needs_rename}")
    print(f"{'='*70}")
    
    # ایجاد گزارش Excel
    excel_path = os.path.join(folder, "گزارش_تصحیح_نام_فایل‌ها.xlsx")
    create_excel_report(results, excel_path)
    
    # تغییر نام؟
    if needs_rename > 0:
        print(f"\n{'='*70}")
        answer = input(f"آیا می‌خواهید {needs_rename} فایل تصحیح شوند؟ (yes/no): ").strip().lower()
        
        if answer in ['yes', 'y', 'بله']:
            rename_files(results)
        else:
            print("\n❌ تصحیح لغو شد. گزارش Excel ذخیره شده است.")
    
    print(f"\n{'='*70}")
    print("✅ پایان برنامه")
    print(f"{'='*70}")

if __name__ == "__main__":
    main()
