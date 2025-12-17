#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import re
from pathlib import Path
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

try:
    import pdfplumber
    HAS_PDF = True
except:
    HAS_PDF = False
    print("WARNING: pdfplumber not installed")

try:
    from docx import Document
    HAS_DOCX = True
except:
    HAS_DOCX = False
    print("WARNING: python-docx not installed")


def extract_from_pdf(pdf_path):
    if not HAS_PDF:
        return None, None, None
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            page = pdf.pages[0]
            text = page.extract_text()
            
            if not text:
                return None, None, None
            
            print(f"    PDF text extracted: {len(text)} chars")
            
            # Find Sequence Number
            seq_num = None
            for pattern in [r'(\d{4})', r'(\d{3})', r'(\d{2})', r'(\d{1})']:
                matches = re.findall(pattern, text)
                if matches:
                    for match in matches:
                        if len(match) >= 1:
                            seq_num = match.zfill(4)
                            break
                    if seq_num:
                        break
            
            # Find Revision
            revision = None
            rev_matches = re.findall(r'(G\d{2,3})', text, re.I)
            if rev_matches:
                revision = rev_matches[0].upper()
            
            # Find Date
            date = None
            date_match = re.search(r'(\d{1,2}[-/][A-Za-z]{3}[-/]\d{4})', text)
            if date_match:
                date = date_match.group(1)
            
            print(f"    Extracted: Seq={seq_num}, Rev={revision}, Date={date}")
            return seq_num, revision, date
            
    except Exception as e:
        print(f"    ERROR reading PDF: {e}")
        return None, None, None


def extract_from_docx(docx_path):
    if not HAS_DOCX:
        return None, None, None
    
    try:
        doc = Document(docx_path)
        
        # Extract ALL text from tables
        text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
            text += "\n"
        
        # Extract ALL text from paragraphs
        for para in doc.paragraphs:
            text += para.text + " "
        
        if not text.strip():
            return None, None, None
        
        print(f"    DOCX text extracted: {len(text)} chars")
        
        # Find Sequence Number - look for any 4-digit or 3-digit number
        seq_num = None
        for pattern in [r'(\d{4})', r'(\d{3})', r'(\d{2})', r'(\d{1})']:
            matches = re.findall(pattern, text)
            if matches:
                for match in matches:
                    if len(match) >= 1 and match != '2024' and match != '2025':
                        seq_num = match.zfill(4)
                        break
                if seq_num:
                    break
        
        # Find Revision
        revision = None
        rev_matches = re.findall(r'(G\d{2,3})', text, re.I)
        if rev_matches:
            revision = rev_matches[0].upper()
        
        # Find Date
        date = None
        date_match = re.search(r'(\d{1,2}[-/][A-Za-z]{3}[-/]\d{4})', text)
        if date_match:
            date = date_match.group(1)
        
        print(f"    Extracted: Seq={seq_num}, Rev={revision}, Date={date}")
        return seq_num, revision, date
        
    except Exception as e:
        print(f"    ERROR reading DOCX: {e}")
        import traceback
        traceback.print_exc()
        return None, None, None


def process_directory(directory_path, output_format):
    directory = Path(directory_path)
    results = []
    failed = []
    
    if not directory.exists():
        print(f"\nERROR: Directory does not exist!")
        print(f"Path: {directory_path}")
        return results, failed
    
    # Find all files
    pdf_files = list(directory.glob("*.pdf"))
    docx_files = list(directory.glob("*.docx"))
    doc_files = list(directory.glob("*.doc"))
    all_files = pdf_files + docx_files + doc_files
    
    print(f"\nFound files:")
    print(f"  PDF: {len(pdf_files)}")
    print(f"  DOCX: {len(docx_files)}")
    print(f"  DOC: {len(doc_files)}")
    print(f"  Total: {len(all_files)}\n")
    
    for file_path in all_files:
        # Skip temp files
        if file_path.name.startswith('~$'):
            print(f"SKIP temp file: {file_path.name}\n")
            continue
        
        print(f"Processing: {file_path.name}")
        print(f"  Type: {file_path.suffix}")
        
        # Extract data
        seq = None
        rev = None
        date = None
        
        if file_path.suffix.lower() == '.pdf':
            seq, rev, date = extract_from_pdf(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            seq, rev, date = extract_from_docx(file_path)
        
        # Check if extraction succeeded
        if seq and rev:
            new_name = f"{output_format}-{seq}-{rev}{file_path.suffix}"
            results.append({
                'original_name': file_path.name,
                'new_name': new_name,
                'date': date or 'N/A',
                'sequence': seq,
                'revision': rev,
                'old_path': file_path
            })
            print(f"  SUCCESS -> {new_name}\n")
        else:
            failed.append({
                'filename': file_path.name,
                'seq': seq,
                'rev': rev,
                'date': date
            })
            print(f"  FAILED - Missing data (Seq={seq}, Rev={rev})\n")
    
    return results, failed


def create_excel_report(results, failed, output_path):
    print(f"\nCreating Excel report...")
    
    try:
        wb = openpyxl.Workbook()
        
        # Sheet 1: Successful
        ws1 = wb.active
        ws1.title = "Successfully Processed"
        
        headers = ['Original Filename', 'New Filename', 'Date', 'Sequence', 'Revision']
        ws1.append(headers)
        
        # Header style
        header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        header_font = Font(bold=True, color='FFFFFF', size=12)
        
        for cell in ws1[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # Add data
        for item in results:
            ws1.append([
                item['original_name'],
                item['new_name'],
                item['date'],
                item['sequence'],
                item['revision']
            ])
        
        # Column widths
        ws1.column_dimensions['A'].width = 50
        ws1.column_dimensions['B'].width = 45
        ws1.column_dimensions['C'].width = 15
        ws1.column_dimensions['D'].width = 12
        ws1.column_dimensions['E'].width = 12
        
        # Sheet 2: Failed
        if failed:
            ws2 = wb.create_sheet("Failed")
            headers2 = ['Filename', 'Sequence Found', 'Revision Found', 'Date Found']
            ws2.append(headers2)
            
            for cell in ws2[1]:
                cell.fill = PatternFill(start_color='FF0000', end_color='FF0000', fill_type='solid')
                cell.font = Font(bold=True, color='FFFFFF', size=12)
                cell.alignment = Alignment(horizontal='center', vertical='center')
            
            for item in failed:
                ws2.append([
                    item['filename'],
                    item['seq'] or 'NOT FOUND',
                    item['rev'] or 'NOT FOUND',
                    item['date'] or 'NOT FOUND'
                ])
            
            ws2.column_dimensions['A'].width = 50
            ws2.column_dimensions['B'].width = 20
            ws2.column_dimensions['C'].width = 20
            ws2.column_dimensions['D'].width = 20
        
        # Save
        wb.save(output_path)
        print(f"SUCCESS: Excel created at {output_path}")
        return True
        
    except Exception as e:
        print(f"ERROR creating Excel: {e}")
        import traceback
        traceback.print_exc()
        return False


def rename_files(results):
    print("\n" + "="*70)
    print("RENAMING FILES")
    print("="*70 + "\n")
    
    success = 0
    skipped = 0
    failed = 0
    
    for item in results:
        old_path = item['old_path']
        new_path = old_path.parent / item['new_name']
        
        try:
            if new_path.exists():
                print(f"SKIP (already exists): {item['new_name']}")
                skipped += 1
                continue
            
            old_path.rename(new_path)
            print(f"OK: {item['original_name']}")
            print(f"    -> {item['new_name']}")
            success += 1
            
        except Exception as e:
            print(f"FAIL: {item['original_name']}")
            print(f"      Error: {e}")
            failed += 1
    
    print("\n" + "="*70)
    print(f"Results: {success} renamed, {skipped} skipped, {failed} failed")
    print("="*70)


def main():
    print("="*70)
    print("FILE RENAMING TOOL")
    print("="*70)
    
    # Choose report type
    print("\nSelect report type:")
    print("1. Monthly Reports (SJSC-GGNRSP-MADR-REMO)")
    print("2. Weekly Reports (SJSC-GGNRSP-EPWC-REWK)")
    
    choice = input("\nEnter choice (1 or 2): ").strip()
    
    if choice == '1':
        DIRECTORY = r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\monthly"
        EXCEL_OUTPUT = r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\monthly_files.xlsx"
        OUTPUT_FORMAT = "SJSC-GGNRSP-MADR-REMO"
        print("\nProcessing MONTHLY reports...")
    elif choice == '2':
        DIRECTORY = r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\weekly"
        EXCEL_OUTPUT = r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\weekly_files.xlsx"
        OUTPUT_FORMAT = "SJSC-GGNRSP-EPWC-REWK"
        print("\nProcessing WEEKLY reports...")
    else:
        print("Invalid choice!")
        return
    
    # Check libraries
    if not HAS_PDF:
        print("\nWARNING: pdfplumber not installed - PDF files will be skipped")
    if not HAS_DOCX:
        print("\nWARNING: python-docx not installed - Word files will be skipped")
    
    if not HAS_PDF and not HAS_DOCX:
        print("\nERROR: No libraries installed!")
        print("Install: pip install pdfplumber python-docx openpyxl")
        return
    
    # Process files
    results, failed = process_directory(DIRECTORY, OUTPUT_FORMAT)
    
    print("\n" + "="*70)
    print(f"SUMMARY:")
    print(f"  Successful: {len(results)}")
    print(f"  Failed: {len(failed)}")
    print("="*70)
    
    # Always create Excel
    excel_created = create_excel_report(results, failed, EXCEL_OUTPUT)
    
    if not results:
        print("\nNo files were successfully processed!")
        print("Check the Excel file for details on failed files.")
        return
    
    # Ask to rename
    print(f"\n{len(results)} files ready to rename")
    answer = input("\nRename files now? (yes/no): ").strip().lower()
    
    if answer in ['yes', 'y']:
        rename_files(results)
    else:
        print("\nFiles NOT renamed. Check Excel file for preview.")
    
    print("\nDONE!")


if __name__ == "__main__":
    main()