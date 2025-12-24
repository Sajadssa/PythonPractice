#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
اسکریپت تغییر نام فایل‌های گزارش ماهانه عملیات تولید
استخراج اطلاعات از جداول درون فایل‌ها و تولید گزارش Excel
برای فایل‌های Word و PDF
"""

import os
import re
from pathlib import Path
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from datetime import datetime

try:
    from PyPDF2 import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("⚠️ توجه: PyPDF2 نصب نیست. قابلیت استخراج از PDF محدود می‌شود.")


class MonthlyFileRenamer:
    def __init__(self, source_directory):
        """
        مقداردهی اولیه
        
        Args:
            source_directory: مسیر پوشه حاوی فایل‌های Word و PDF
        """
        self.source_directory = Path(source_directory)
        self.results = []
        
    def extract_table_data_from_word(self, doc):
        """
        استخراج Sequence Number و Date از جداول فایل Word (از Header و Body)
        
        Args:
            doc: شیء Document از python-docx
            
        Returns:
            dict: دیکشنری حاوی sequence_number, date, revision
        """
        data = {
            'sequence_number': None,
            'date': None,
            'revision': None
        }
        
        # 1. جستجو در Header (برای Sequence Number و Revision)
        for section in doc.sections:
            header = section.header
            for table in header.tables:
                for row_idx, row in enumerate(table.rows):
                    cells_text = [cell.text.strip() for cell in row.cells]
                    
                    # جستجوی Sequence Number
                    for i, cell_text in enumerate(cells_text):
                        if 'Sequence' in cell_text and 'Number' in cell_text:
                            try:
                                if row_idx + 1 < len(table.rows):
                                    next_row = table.rows[row_idx + 1]
                                    seq_value = next_row.cells[i].text.strip()
                                    if seq_value and seq_value.isdigit():
                                        data['sequence_number'] = seq_value.zfill(4)
                            except:
                                pass
                    
                    # جستجوی Revision در Header
                    for i, cell_text in enumerate(cells_text):
                        if cell_text in ['Revision', 'Rev', 'REV']:
                            try:
                                if row_idx + 1 < len(table.rows):
                                    next_row = table.rows[row_idx + 1]
                                    rev_value = next_row.cells[i].text.strip()
                                    if rev_value and rev_value.startswith('G'):
                                        data['revision'] = rev_value
                            except:
                                pass
        
        # 2. جستجو در جداول اصلی سند (برای Date و Revision)
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                cells_text = [cell.text.strip() for cell in row.cells]
                
                # جستجوی Revision (اگر در Header پیدا نشده)
                if not data['revision']:
                    for i, cell_text in enumerate(cells_text):
                        if cell_text in ['Revision', 'Rev', 'REV']:
                            try:
                                if row_idx + 1 < len(table.rows):
                                    next_row = table.rows[row_idx + 1]
                                    rev_value = next_row.cells[i].text.strip()
                                    if rev_value:
                                        data['revision'] = rev_value
                            except:
                                pass
                
                # جستجوی Date
                for i, cell_text in enumerate(cells_text):
                    if cell_text == 'Date' or cell_text == 'DATE':
                        try:
                            if row_idx + 1 < len(table.rows):
                                next_row = table.rows[row_idx + 1]
                                date_value = next_row.cells[i].text.strip()
                                if date_value and len(date_value) > 5:
                                    data['date'] = date_value
                        except:
                            pass
        
        return data
    
    def extract_table_data_from_pdf(self, file_path):
        """
        استخراج اطلاعات از فایل PDF
        
        Args:
            file_path: مسیر فایل PDF
            
        Returns:
            dict: دیکشنری حاوی sequence_number, date, revision
        """
        data = {
            'sequence_number': None,
            'date': None,
            'revision': None
        }
        
        if not PDF_SUPPORT:
            return data
        
        try:
            reader = PdfReader(file_path)
            
            # استخراج متن از صفحه اول
            if len(reader.pages) > 0:
                first_page = reader.pages[0]
                text = first_page.extract_text()
                
                # جستجوی Sequence Number - الگوهای مختلف
                # الگو 1: "Sequence Number 0003"
                seq_patterns = [
                    r'Sequence\s+Number\s+(\d{4})',
                    r'Sequence\s+Number\s*(\d{3,4})',
                    r'REWK\s+(\d{4})',
                ]
                
                for pattern in seq_patterns:
                    match = re.search(pattern, text, re.IGNORECASE)
                    if match:
                        data['sequence_number'] = match.group(1).zfill(4)
                        break
                
                # جستجوی Revision - الگوهای مختلف
                rev_patterns = [
                    r'Revision\s+(G\d{2})',
                    r'\b(G\d{2})\b',
                ]
                
                for pattern in rev_patterns:
                    match = re.search(pattern, text)
                    if match:
                        data['revision'] = match.group(1)
                        break
                
                # جستجوی Date - الگوهای مختلف تاریخ
                date_patterns = [
                    r'\b(\d{1,2}-[A-Za-z]{3}-\d{4})\b',  # 05-Feb-2024
                    r'\b(\d{1,2}-[A-Za-z]{3}-\d{2})\b',   # 27-Jan-24
                ]
                
                for pattern in date_patterns:
                    dates = re.findall(pattern, text)
                    if dates:
                        data['date'] = dates[0]
                        break
            
        except Exception as e:
            print(f"      خطا در خواندن PDF: {str(e)}")
        
        return data
    
    def generate_new_filename(self, sequence_number, revision, extension):
        """
        تولید نام جدید فایل برای Monthly
        
        Args:
            sequence_number: شماره توالی
            revision: شماره بازنگری
            extension: پسوند فایل (.docx یا .pdf)
            
        Returns:
            str: نام جدید فایل
        """
        # فرمت: SJSC-GGNRSP-EPWC-REMO-[Sequence]-[Revision]
        if not sequence_number:
            sequence_number = "0000"
        if not revision:
            revision = "G00"
        
        # اگر revision شامل G نباشد، آن را اضافه کن
        if not revision.startswith('G'):
            revision = 'G' + revision
            # SJSC-GGNRSP-EPME-REWK-0001-G00
        new_name = f"SJSC-GGNRSP-EPME-REWK-{sequence_number}-{revision}{extension}"
        return new_name
    
    def process_files(self, rename_files=False):
        """
        پردازش فایل‌های Word و PDF
        
        Args:
            rename_files: آیا فایل‌ها تغییر نام پیدا کنند؟
            
        Returns:
            list: لیست نتایج
        """
        # پیدا کردن تمام فایل‌های Word
        word_files = list(self.source_directory.glob("*.docx"))
        word_files.extend(list(self.source_directory.glob("*.doc")))
        
        # پیدا کردن تمام فایل‌های PDF
        pdf_files = list(self.source_directory.glob("*.pdf"))
        
        print(f"تعداد {len(word_files)} فایل Word پیدا شد.")
        print(f"تعداد {len(pdf_files)} فایل PDF پیدا شد.\n")
        
        # پردازش فایل‌های Word
        for file_path in word_files:
            # رد کردن فایل‌های موقت
            if file_path.name.startswith('~$'):
                continue
                
            print(f"📄 در حال پردازش Word: {file_path.name}")
            
            try:
                # خواندن فایل Word
                doc = Document(file_path)
                
                # استخراج داده‌ها
                data = self.extract_table_data_from_word(doc)
                
                print(f"  Sequence Number: {data['sequence_number']}")
                print(f"  Revision: {data['revision']}")
                print(f"  Date: {data['date']}")
                
                # تولید نام جدید
                new_filename = self.generate_new_filename(
                    data['sequence_number'], 
                    data['revision'],
                    '.docx'
                )
                
                print(f"  نام جدید: {new_filename}\n")
                
                # ذخیره نتیجه
                result = {
                    'file_type': 'Word',
                    'original_name': file_path.name,
                    'new_name': new_filename,
                    'sequence_number': data['sequence_number'] or 'N/A',
                    'revision': data['revision'] or 'N/A',
                    'date': data['date'] or 'N/A',
                    'status': 'پردازش شد'
                }
                
                # تغییر نام فایل (اختیاری)
                if rename_files:
                    new_path = file_path.parent / new_filename
                    if not new_path.exists():
                        file_path.rename(new_path)
                        result['status'] = 'تغییر نام داده شد'
                    else:
                        result['status'] = 'فایل با این نام وجود دارد'
                
                self.results.append(result)
                
            except Exception as e:
                print(f"  خطا: {str(e)}\n")
                self.results.append({
                    'file_type': 'Word',
                    'original_name': file_path.name,
                    'new_name': 'خطا',
                    'sequence_number': 'N/A',
                    'revision': 'N/A',
                    'date': 'N/A',
                    'status': f'خطا: {str(e)}'
                })
        
        # پردازش فایل‌های PDF
        for file_path in pdf_files:
            print(f"📕 در حال پردازش PDF: {file_path.name}")
            
            try:
                # استخراج داده‌ها از PDF
                data = self.extract_table_data_from_pdf(file_path)
                
                print(f"  Sequence Number: {data['sequence_number']}")
                print(f"  Revision: {data['revision']}")
                print(f"  Date: {data['date']}")
                
                # تولید نام جدید
                new_filename = self.generate_new_filename(
                    data['sequence_number'], 
                    data['revision'],
                    '.pdf'
                )
                
                print(f"  نام جدید: {new_filename}\n")
                
                # ذخیره نتیجه
                result = {
                    'file_type': 'PDF',
                    'original_name': file_path.name,
                    'new_name': new_filename,
                    'sequence_number': data['sequence_number'] or 'N/A',
                    'revision': data['revision'] or 'N/A',
                    'date': data['date'] or 'N/A',
                    'status': 'پردازش شد'
                }
                
                # تغییر نام فایل (اختیاری)
                if rename_files:
                    new_path = file_path.parent / new_filename
                    if not new_path.exists():
                        file_path.rename(new_path)
                        result['status'] = 'تغییر نام داده شد'
                    else:
                        result['status'] = 'فایل با این نام وجود دارد'
                
                self.results.append(result)
                
            except Exception as e:
                print(f"  خطا: {str(e)}\n")
                self.results.append({
                    'file_type': 'PDF',
                    'original_name': file_path.name,
                    'new_name': 'خطا',
                    'sequence_number': 'N/A',
                    'revision': 'N/A',
                    'date': 'N/A',
                    'status': f'خطا: {str(e)}'
                })
        
        return self.results
    
    def create_excel_report(self, output_path):
        """
        تولید گزارش Excel
        
        Args:
            output_path: مسیر فایل خروجی Excel
        """
        wb = Workbook()
        ws = wb.active
        ws.title = "گزارش فایل‌های ماهانه"
        
        # تنظیمات سبک
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=12)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # هدر جدول
        headers = ['ردیف', 'نوع فایل', 'نام اصلی فایل', 'نام جدید فایل', 'Sequence Number', 'Revision', 'Date', 'وضعیت']
        
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border
        
        # داده‌ها
        for row_idx, result in enumerate(self.results, 2):
            ws.cell(row=row_idx, column=1, value=row_idx-1).border = border
            ws.cell(row=row_idx, column=2, value=result.get('file_type', 'N/A')).border = border
            ws.cell(row=row_idx, column=3, value=result['original_name']).border = border
            ws.cell(row=row_idx, column=4, value=result['new_name']).border = border
            ws.cell(row=row_idx, column=5, value=result['sequence_number']).border = border
            ws.cell(row=row_idx, column=6, value=result['revision']).border = border
            ws.cell(row=row_idx, column=7, value=result['date']).border = border
            ws.cell(row=row_idx, column=8, value=result['status']).border = border
        
        # تنظیم عرض ستون‌ها
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 12
        ws.column_dimensions['C'].width = 50
        ws.column_dimensions['D'].width = 50
        ws.column_dimensions['E'].width = 18
        ws.column_dimensions['F'].width = 12
        ws.column_dimensions['G'].width = 15
        ws.column_dimensions['H'].width = 20
        
        # ذخیره فایل
        wb.save(output_path)
        print(f"\n✓ گزارش Excel در مسیر زیر ذخیره شد:")
        print(f"  {output_path}")


def main():
    """
    تابع اصلی برنامه
    """
    print("=" * 70)
    print("برنامه تغییر نام فایل‌های گزارش ماهانه عملیات تولید")
    print("=" * 70)
    print()
    
    # مسیر پوشه فایل‌ها
    source_dir = r"D:\Sepher_Pasargad\works\Production\Corrosion\Weekly"
    
    # بررسی وجود پوشه
    if not os.path.exists(source_dir):
        print(f"خطا: پوشه {source_dir} یافت نشد!")
        print("لطفاً مسیر را در کد بررسی کنید.")
        return
    
    # ایجاد شیء پردازشگر
    renamer = MonthlyFileRenamer(source_dir)
    
    # پردازش فایل‌ها (بدون تغییر نام واقعی)
    print("در حال پردازش فایل‌ها...")
    print("-" * 70)
    results = renamer.process_files(rename_files=False)
    
    # تولید گزارش Excel
    output_excel = os.path.join(source_dir, f"گزارش_ماهانه_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
    renamer.create_excel_report(output_excel)
    
    # خلاصه نتایج
    print("\n" + "=" * 70)
    print("خلاصه نتایج:")
    print(f"  تعداد کل فایل‌ها: {len(results)}")
    print(f"  Word: {sum(1 for r in results if r.get('file_type') == 'Word')}")
    print(f"  PDF: {sum(1 for r in results if r.get('file_type') == 'PDF')}")
    print(f"  موفق: {sum(1 for r in results if 'خطا' not in r['status'])}")
    print(f"  خطا: {sum(1 for r in results if 'خطا' in r['status'])}")
    print("=" * 70)
    
    # سوال برای تغییر نام واقعی
    print("\nآیا می‌خواهید فایل‌ها واقعاً تغییر نام پیدا کنند?")
    print("توجه: این عملیات قابل بازگشت نیست!")
    choice = input("برای تغییر نام 'yes' وارد کنید: ")
    
    if choice.lower() == 'yes':
        print("\nدر حال تغییر نام فایل‌ها...")
        renamer.results = []
        results = renamer.process_files(rename_files=True)
        
        # تولید گزارش جدید
        output_excel_final = os.path.join(source_dir, f"گزارش_نهایی_ماهانه_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
        renamer.create_excel_report(output_excel_final)
        print("\n✓ فایل‌ها با موفقیت تغییر نام یافتند!")


if __name__ == "__main__":
    main()
