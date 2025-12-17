#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import re
from pathlib import Path
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

try:
    import pdfplumber
except:
    pdfplumber = None

try:
    from docx import Document
except:
    Document = None


def extract_from_pdf(pdf_path):
    if not pdfplumber:
        return None, None, None
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = pdf.pages[0].extract_text()
            
            # Sequence Number
            seq_num = None
            patterns = [
                r'Sequence\s+Number\s+(\d+)',
                r'Number\s+(\d+)',
                r'REWK\s+(\d+)',
                r'\|\s*(\d{4})\s*\|',
            ]
            for pattern in patterns:
                seq = re.search(pattern, text, re.I)
                if seq:
                    seq_num = seq.group(1).zfill(4)
                    break
            
            # Revision
            revision = None
            patterns = [
                r'Revision\s+(G\d+)',
                r'\b(G\d{2,3})\b',
                r'G\d{2,3}',
            ]
            for pattern in patterns:
                rev = re.search(pattern, text, re.I)
                if rev:
                    revision = rev.group(1).upper() if '(' in pattern else rev.group(0).upper()
                    break
            
            # Date
            date = re.search(r'(\d{1,2}[-/][A-Za-z]{3}[-/]\d{4})', text)
            date_str = date.group(1) if date else None
            
            return seq_num, revision, date_str
    except Exception as e:
        print(f"    Error: {e}")
        return None, None, None


def extract_from_docx(docx_path):
    if not Document:
        return None, None, None
    
    try:
        doc = Document(docx_path)
        text = ""
        
        for table in doc.tables[:2]:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
        
        for para in doc.paragraphs[:10]:
            text += para.text + " "
        
        # Sequence Number
        seq_num = None
        patterns = [
            r'Sequence\s+Number\s+(\d+)',
            r'Number\s+(\d+)',
            r'REWK\s+(\d+)',
            r'\|\s*(\d{4})\s*\|',
        ]
        for pattern in patterns:
            seq = re.search(pattern, text, re.I)
            if seq:
                seq_num = seq.group(1).zfill(4)
                break
        
        # Revision
        revision = None
        patterns = [
            r'Revision\s+(G\d+)',
            r'\b(G\d{2,3})\b',
            r'G\d{2,3}',
        ]
        for pattern in patterns:
            rev = re.search(pattern, text, re.I)
            if rev:
                revision = rev.group(1).upper() if '(' in pattern else rev.group(0).upper()
                break
        
        # Date
        date = re.search(r'(\d{1,2}[-/][A-Za-z]{3}[-/]\d{4})', text)
        date_str = date.group(1) if date else None
        
        return seq_num, revision, date_str
    except Exception as e:
        print(f"    Error: {e}")
        return None, None, None


def process_directory(directory_path):
    directory = Path(directory_path)
    results = []
    
    if not directory.exists():
        print(f"ERROR: Directory not found: {directory_path}")
        return results
    
    all_files = list(directory.glob("*.pdf")) + list(directory.glob("*.docx")) + list(directory.glob("*.doc"))
    
    print(f"\nFound {len(all_files)} files\n")
    
    for file_path in all_files:
        if file_path.name.startswith('~$'):
            continue
        
        print(f"Processing: {file_path.name}")
        
        if file_path.suffix.lower() == '.pdf':
            seq, rev, date = extract_from_pdf(file_path)
        else:
            seq, rev, date = extract_from_docx(file_path)
        
        if seq and rev:
            new_name = f"SJSC-GGNRSP-EPWC-REWK-{seq}-{rev}{file_path.suffix}"
            results.append({
                'original_name': file_path.name,
                'new_name': new_name,
                'date': date or 'N/A',
                'old_path': file_path
            })
            print(f"  -> {new_name}")
            print(f"  Date: {date}\n")
        else:
            print(f"  FAILED - Seq: {seq}, Rev: {rev}\n")
    
    return results


def create_excel_report(results, output_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Renamed Files"
    
    headers = ['Original Filename', 'New Filename', 'Date']
    ws.append(headers)
    
    header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF', size=12)
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    for item in results:
        ws.append([
            item['original_name'],
            item['new_name'],
            item['date']
        ])
    
    ws.column_dimensions['A'].width = 50
    ws.column_dimensions['B'].width = 45
    ws.column_dimensions['C'].width = 20
    
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        for cell in row:
            cell.alignment = Alignment(horizontal='left', vertical='center')
        row[2].alignment = Alignment(horizontal='center', vertical='center')
    
    wb.save(output_path)
    print(f"\nExcel report created: {output_path}\n")


def rename_files(results):
    print("\nRenaming files...\n")
    success = 0
    
    for item in results:
        old_path = item['old_path']
        new_path = old_path.parent / item['new_name']
        
        if new_path.exists():
            print(f"SKIP (exists): {item['new_name']}")
            continue
        
        try:
            old_path.rename(new_path)
            print(f"OK: {item['original_name']} -> {item['new_name']}")
            success += 1
        except Exception as e:
            print(f"FAIL: {item['original_name']} - {e}")
    
    print(f"\nRenamed {success}/{len(results)} files")


def main():
    DIRECTORY = r"D:\Sepher_Pasargad\works\Production\Production Operation Report\Weekly"
    EXCEL_OUTPUT = r"D:\Sepher_Pasargad\works\Production\Production Operation Report\weekly_files.xlsx"
    
    print("="*70)
    print("WEEKLY REPORT RENAMING TOOL")
    print("Format: SJSC-GGNRSP-EPWC-REWK-####-G##")
    print("="*70)
    
    results = process_directory(DIRECTORY)
    
    if not results:
        print("\nNo files processed!")
        return
    
    print(f"\n{len(results)} files ready to rename")
    
    try:
        create_excel_report(results, EXCEL_OUTPUT)
    except Exception as e:
        print(f"ERROR creating Excel: {e}")
    
    answer = input("\nRename files now? (yes/no): ").strip().lower()
    
    if answer in ['yes', 'y']:
        rename_files(results)
        print("\nDone!")
    else:
        print("\nFiles not renamed. Check Excel file.")


if __name__ == "__main__":
    main()
