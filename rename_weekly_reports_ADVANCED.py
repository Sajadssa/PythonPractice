#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
تغییر نام Maintenance Weekly Reports
نسخه نهایی - با رفع کامل باگ استخراج Sequence Number
استفاده از Chain of Thought (CoT) برای دیباگ
"""

import os
from pathlib import Path
import re
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from collections import defaultdict

# سعی در import کتابخانه‌های PDF
try:
    import PyPDF2
    PDF_SUPPORT = True
except:
    PDF_SUPPORT = False
    print("⚠️ PyPDF2 نصب نیست - فقط فایل‌های Word پردازش می‌شوند")

try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_AVAILABLE = True
except:
    OCR_AVAILABLE = False

def parse_date(date_str):
    """تبدیل تاریخ به datetime"""
    if not date_str:
        return None
    
    months = {
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
    }
    
    try:
        date_str = date_str.strip()
        parts = date_str.split('-')
        if len(parts) == 3:
            day = int(parts[0])
            month = months.get(parts[1].lower())
            year = int(parts[2])
            if month:
                return datetime(year, month, day)
    except:
        pass
    return None

def extract_from_word(doc_path):
    """
    استخراج اطلاعات از فایل Word
    
    🎯 روش صحیح: استفاده از نام ستون‌ها به جای جستجوی کور
    
    مشکل قدیمی: کد اولین عدد 4 رقمی که پیدا می‌کرد را برمی‌گرداند
    راه‌حل جدید: با استفاده از نام ستون "Sequence Number" مقدار صحیح را استخراج می‌کند
    """
    print(f"   📄 پردازش Word...")
    
    try:
        doc = Document(doc_path)
        print(f"   📊 جداول: {len(doc.tables)}")
        
        sequence_number = None
        revision = None
        date_obj = None
        date_str = None
        
        # 🎯 استراتژی 1: استفاده از نام ستون‌ها (بهترین روش)
        for table_idx, table in enumerate(doc.tables):
            if sequence_number and revision and date_obj:
                break
            
            if len(table.rows) < 2:
                continue
            
            # خواندن ردیف هدر
            headers = []
            for cell in table.rows[0].cells:
                # حذف newline و فضاهای اضافی
                header = cell.text.strip().replace('\n', ' ').replace('\r', ' ')
                headers.append(header)
            
            print(f"   📋 جدول {table_idx+1} - Headers: {headers}")
            
            # پیدا کردن index ستون‌های مورد نظر
            seq_idx = -1
            rev_idx = -1
            date_idx = -1
            
            for idx, header in enumerate(headers):
                header_lower = header.lower()
                
                # جستجوی ستون Sequence Number
                if 'sequence' in header_lower and 'number' in header_lower:
                    seq_idx = idx
                elif 'sequence' in header_lower or header.lower() == 'seq':
                    seq_idx = idx
                
                # جستجوی ستون Revision
                if 'revision' in header_lower or header.lower() == 'rev':
                    rev_idx = idx
                
                # جستجوی ستون Date
                if 'date' in header_lower:
                    date_idx = idx
            
            print(f"   📌 Indices: Seq={seq_idx}, Rev={rev_idx}, Date={date_idx}")
            
            # استخراج داده‌ها از ستون‌های شناسایی شده
            # بررسی تمام ردیف‌های داده (نه فقط ردیف 2)
            for row_idx in range(1, len(table.rows)):
                data_row = table.rows[row_idx]
                
                # Sequence Number
                if not sequence_number and seq_idx >= 0 and seq_idx < len(data_row.cells):
                    val = data_row.cells[seq_idx].text.strip()
                    # فقط اگر عدد باشد
                    if val and (val.isdigit() or (len(val) == 4 and val.replace('0', '').isdigit())):
                        sequence_number = val
                        print(f"   ✅ Sequence: {val} (جدول {table_idx+1}, ستون {seq_idx+1}, ردیف {row_idx+1})")
                
                # Revision
                if not revision and rev_idx >= 0 and rev_idx < len(data_row.cells):
                    val = data_row.cells[rev_idx].text.strip()
                    # الگوی Revision: G + 2 رقم
                    if val and len(val) >= 2 and val[0].upper() == 'G' and val[1:].isdigit():
                        revision = val.upper()
                        print(f"   ✅ Revision: {val} (جدول {table_idx+1}, ستون {rev_idx+1}, ردیف {row_idx+1})")
                
                # Date
                if not date_obj and date_idx >= 0 and date_idx < len(data_row.cells):
                    val = data_row.cells[date_idx].text.strip()
                    temp_date = parse_date(val)
                    if temp_date:
                        date_obj = temp_date
                        date_str = val
                        print(f"   ✅ Date: {val} (جدول {table_idx+1}, ستون {date_idx+1}, ردیف {row_idx+1})")
                
                if sequence_number and revision and date_obj:
                    break
        
        # 🎯 استراتژی 2: اگر از هدرها نتوانستیم پیدا کنیم، جستجوی هوشمند
        if not sequence_number or not revision:
            print(f"   🔄 استراتژی 2: جستجوی هوشمند در جداول...")
            
            for table_idx, table in enumerate(doc.tables):
                if sequence_number and revision:
                    break
                
                if len(table.rows) < 2:
                    continue
                
                # بررسی تمام ردیف‌ها
                for row_idx in range(1, len(table.rows)):
                    row = table.rows[row_idx]
                    
                    # استخراج مقادیر
                    values = []
                    for cell in row.cells:
                        val = cell.text.strip()
                        if val:
                            values.append(val)
                    
                    # جستجوی Revision (اولویت دارد چون منحصر به فردتر است)
                    if not revision:
                        for val in values:
                            if len(val) == 3 and val[0].upper() == 'G' and val[1:].isdigit():
                                revision = val.upper()
                                print(f"   ✅ Revision (جستجو): {val}")
                                break
                    
                    # جستجوی Sequence: فقط اعداد 4 رقمی که با 0 شروع می‌شوند یا کوچکتر از 2000 هستند
                    # (چون معمولاً Sequence Number‌ها اعداد کوچکی هستند)
                    if not sequence_number:
                        for val in values:
                            if val.isdigit() and len(val) == 4:
                                # بررسی: آیا این Sequence محتمل است یا شماره سند؟
                                # Sequence معمولاً با 0 شروع می‌شود یا کوچک است
                                num = int(val)
                                if val.startswith('0') or num < 2000:
                                    sequence_number = val
                                    print(f"   ✅ Sequence (جستجو): {val}")
                                    break
        
        # 🎯 استراتژی 3: استفاده از Regex فقط در صورت عدم موفقیت روش‌های قبل
        if not sequence_number or not revision:
            print(f"   🔄 استراتژی 3: جستجوی Regex...")
            
            all_text = ""
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text += cell.text + " "
            
            for para in doc.paragraphs:
                all_text += para.text + " "
            
            # الگوهای خاص برای REWK documents
            patterns = [
                r'REWK\s+(\d{4})\s+(G\d{2})',      # REWK 0048 G00
                r'REWK[|\s]+(\d{4})[|\s]+(G\d{2})', # REWK|0048|G00
            ]
            
            for pattern in patterns:
                match = re.search(pattern, all_text)
                if match:
                    if not sequence_number:
                        # گرفتن اعدادی که با 0 شروع می‌شوند
                        potential_seq = match.group(1)
                        if potential_seq.startswith('0') or int(potential_seq) < 2000:
                            sequence_number = potential_seq
                    if not revision:
                        revision = match.group(2).upper()
                    if sequence_number and revision:
                        print(f"   ✅ از Regex: {sequence_number}-{revision}")
                        break
        
        # نتیجه نهایی
        if sequence_number and revision:
            print(f"   ✅ نتیجه نهایی: {sequence_number}-{revision}")
            return {
                'sequence_number': sequence_number,
                'revision': revision,
                'date': date_obj,
                'date_str': date_str
            }
        else:
            print(f"   ❌ استخراج ناموفق")
            print(f"      Sequence: {sequence_number}")
            print(f"      Revision: {revision}")
            return None
            
    except Exception as e:
        print(f"   ❌ خطا: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def extract_from_pdf(pdf_path):
    """استخراج اطلاعات از فایل PDF"""
    if not PDF_SUPPORT:
        print(f"   ⚠️ PyPDF2 نصب نیست")
        return None
    
    print(f"   📄 پردازش PDF...")
    
    try:
        text = ""
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            if len(reader.pages) > 0:
                text = reader.pages[0].extract_text()
        
        if not text or len(text) < 50:
            if OCR_AVAILABLE:
                print(f"   🔍 استفاده از OCR...")
                images = convert_from_path(pdf_path, first_page=1, last_page=1, dpi=300)
                if images:
                    text = pytesseract.image_to_string(images[0], lang='eng')
        
        if not text:
            return None
        
        print(f"   📄 متن: {len(text)} کاراکتر")
        
        # استخراج با Regex
        seq_rev_match = re.search(r'REWK\s+(\d{4})\s+(G\d{2})', text)
        if not seq_rev_match:
            seq_rev_match = re.search(r'(\d{4})\s+(G\d{2})', text)
        
        date_match = re.search(r'(\d{1,2}-[A-Za-z]{3,9}-\d{4})', text)
        
        if seq_rev_match:
            sequence_number = seq_rev_match.group(1)
            revision = seq_rev_match.group(2)
            date_str = date_match.group(1) if date_match else None
            date_obj = parse_date(date_str) if date_str else None
            
            print(f"   ✅ نتیجه: {sequence_number}-{revision}")
            
            return {
                'sequence_number': sequence_number,
                'revision': revision,
                'date': date_obj,
                'date_str': date_str
            }
        else:
            print(f"   ❌ الگو پیدا نشد")
            return None
            
    except Exception as e:
        print(f"   ❌ خطا: {str(e)}")
        return None

def create_excel_report(files_data, output_path):
    """ایجاد گزارش Excel از نتایج پردازش"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Weekly Reports"
    
    # استایل‌ها
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    
    # هدر
    headers = ['ردیف', 'نام فایل اصلی', 'نام فایل جدید', 'Sequence', 'Revision', 'تاریخ', 'وضعیت']
    ws.append(headers)
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border
    
    # داده‌ها
    for idx, data in enumerate(files_data, start=1):
        row = [
            idx,
            data['old_name'],
            data['new_name'] if data['new_name'] else 'N/A',
            data['sequence_number'] if data['sequence_number'] else 'N/A',
            data['revision'] if data['revision'] else 'N/A',
            data['date'] if data['date'] else 'N/A',
            data['status']
        ]
        ws.append(row)
        
        row_num = idx + 1
        for col_idx, cell in enumerate(ws[row_num], start=1):
            cell.border = border
            cell.alignment = Alignment(horizontal='center', vertical='center')
            if col_idx == 6 and isinstance(cell.value, datetime):
                cell.number_format = 'DD/MM/YYYY'
    
    # تنظیم عرض ستون‌ها
    widths = [8, 40, 45, 15, 12, 15, 20]
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + idx)].width = width
    
    wb.save(output_path)
    print(f"\n📊 گزارش Excel: {output_path.name}")

def main():
    """تابع اصلی برنامه"""
    
    # ⚠️ مهم: مسیر پوشه را اینجا تنظیم کنید
    FOLDER = Path(r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\weekly")
    
    # برای تست در لینوکس:
    # FOLDER = Path('/home/claude')
    
    if not FOLDER.exists():
        print(f"❌ پوشه پیدا نشد: {FOLDER}")
        print("\n💡 راهنمایی:")
        print("   1. مسیر را در خط 415 کد چک کنید")
        print("   2. مطمئن شوید پوشه وجود دارد")
        return
    
    print("\n" + "="*80)
    print(" 🔄 تغییر نام Maintenance Weekly Reports - نسخه پیشرفته")
    print("="*80)
    print(f"📂 مسیر: {FOLDER}\n")
    
    # جمع‌آوری فایل‌ها
    pdf_files = list(FOLDER.glob('*.pdf'))
    word_files = list(FOLDER.glob('*.docx')) + list(FOLDER.glob('*.doc'))
    
    # فیلتر کردن فایل‌هایی که قبلاً rename شده‌اند
    pdf_files = [f for f in pdf_files if not f.name.startswith('SJSC-GGNRSP-MADR-REWK-')]
    word_files = [f for f in word_files 
                  if not f.name.startswith('SJSC-GGNRSP-MADR-REWK-')
                  and not f.name.startswith('~')
                  and not f.name.startswith('.')]
    
    all_files = pdf_files + word_files
    
    print(f"📁 فایل‌های پیدا شده:")
    print(f"   PDF: {len(pdf_files)}")
    print(f"   Word: {len(word_files)}")
    print(f"   جمع: {len(all_files)}\n")
    
    if not all_files:
        print("❌ فایلی برای پردازش پیدا نشد!")
        print("\n💡 احتمالاً:")
        print("   - همه فایل‌ها قبلاً rename شده‌اند")
        print("   - پوشه خالی است")
        return
    
    print("⚠️ هشدار: این برنامه نام فایل‌ها را تغییر می‌دهد.")
    print("⚠️ فرمت جدید: SJSC-GGNRSP-MADR-REWK-[Sequence]-[Revision]")
    print("\n⚠️ لطفاً قبل از ادامه از فایل‌های خود پشتیبان بگیرید!")
    print("\nآیا مطمئن هستید که می‌خواهید ادامه دهید؟ (y/n): ", end='')
    
    response = input().lower()
    if response != 'y':
        print("❌ عملیات لغو شد")
        return
    
    print("\n🔍 شروع پردازش فایل‌ها...")
    print("-"*80)
    
    # پردازش فایل‌ها
    files_data = []
    
    for file_path in all_files:
        print(f"\n📄 {file_path.name}")
        
        # استخراج اطلاعات
        if file_path.suffix.lower() == '.pdf':
            info = extract_from_pdf(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            info = extract_from_word(file_path)
        else:
            info = None
        
        # ذخیره نتیجه
        if info and info['sequence_number'] and info['revision']:
            files_data.append({
                'path': file_path,
                'old_name': file_path.name,
                'sequence_number': info['sequence_number'],
                'revision': info['revision'],
                'date': info['date'],
                'date_str': info['date_str'],
                'new_name': None,
                'status': 'آماده'
            })
        else:
            files_data.append({
                'path': file_path,
                'old_name': file_path.name,
                'sequence_number': None,
                'revision': None,
                'date': None,
                'date_str': None,
                'new_name': None,
                'status': '❌ اطلاعات یافت نشد'
            })
    
    print("\n" + "="*80)
    print("🔢 تخصیص نام‌های جدید...")
    
    # گروه‌بندی بر اساس Sequence-Revision
    groups = defaultdict(list)
    for data in files_data:
        if data['sequence_number'] and data['revision']:
            key = f"{data['sequence_number']}-{data['revision']}"
            groups[key].append(data)
    
    # تخصیص نام جدید
    for key, group in groups.items():
        if len(group) == 1:
            # فقط یک فایل با این Sequence-Revision
            data = group[0]
            ext = data['path'].suffix.lower()
            data['new_name'] = f"SJSC-GGNRSP-MADR-REWK-{data['sequence_number']}-{data['revision']}{ext}"
            print(f"   ✅ {key} → {data['new_name']}")
        else:
            # چند فایل تکراری
            print(f"   ⚠️ تکراری: {key}")
            for idx, data in enumerate(group, start=1):
                ext = data['path'].suffix.lower()
                data['new_name'] = f"SJSC-GGNRSP-MADR-REWK-{data['sequence_number']}-{data['revision']}_copy{idx}{ext}"
                print(f"      → {data['new_name']}")
    
    # تغییر نام فایل‌ها
    print("\n🔄 در حال تغییر نام فایل‌ها...")
    print("-"*80)
    
    success = 0
    failed = 0
    
    for data in files_data:
        if data['new_name']:
            old_path = data['path']
            new_path = old_path.parent / data['new_name']
            
            try:
                if new_path.exists() and new_path != old_path:
                    data['status'] = '❌ نام تکراری وجود دارد'
                    failed += 1
                else:
                    old_path.rename(new_path)
                    data['status'] = '✅ موفق'
                    success += 1
                    print(f"✅ {old_path.name}")
                    print(f"   → {data['new_name']}")
            except Exception as e:
                data['status'] = f'❌ خطا: {str(e)}'
                failed += 1
                print(f"❌ {old_path.name}: {str(e)}")
    
    print("-"*80)
    
    # ایجاد گزارش Excel
    report_path = FOLDER / f"Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    create_excel_report(files_data, report_path)
    
    # نمایش خلاصه نتایج
    print(f"\n📊 خلاصه نتایج:")
    print(f"   ✅ موفق: {success}")
    print(f"   ❌ ناموفق: {failed}")
    
    pdf_success = sum(1 for d in files_data if '.pdf' in d['old_name'] and d['status'] == '✅ موفق')
    word_success = sum(1 for d in files_data if '.doc' in d['old_name'] and d['status'] == '✅ موفق')
    
    print(f"   📄 PDF: {pdf_success}")
    print(f"   📝 Word: {word_success}")
    print("="*80)
    print("\n✨ عملیات تمام شد!")
    print(f"📊 گزارش کامل در فایل Excel ذخیره شد")

if __name__ == "__main__":
    main()
