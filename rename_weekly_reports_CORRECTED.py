#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
تغییر نام Maintenance Weekly Reports
نسخه نهایی اصلاح شده - با استفاده از Header برای شناسایی دقیق ستون‌ها
"""

import os
from pathlib import Path
import re
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from collections import defaultdict

# سعی در import کتابخانه‌های PDF
try:
    import PyPDF2
    PDF_SUPPORT = True
except:
    PDF_SUPPORT = False
    print("⚠️ PyPDF2 نصب نیست - فقط فایل‌های Word پردازش می‌شوند")

try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_AVAILABLE = True
except:
    OCR_AVAILABLE = False

def parse_date(date_str):
    """تبدیل تاریخ به datetime"""
    if not date_str:
        return None
    
    months = {
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
    }
    
    try:
        date_str = date_str.strip()
        parts = date_str.split('-')
        if len(parts) == 3:
            day = int(parts[0])
            month = months.get(parts[1].lower())
            year = int(parts[2])
            if month:
                return datetime(year, month, day)
    except:
        pass
    return None

def extract_from_word(doc_path):
    """
    استخراج اطلاعات از فایل Word با استفاده از Header
    
    این تابع از سه روش استفاده می‌کند (به ترتیب اولویت):
    1. خواندن header و شناسایی ستون‌های Sequence/Revision (دقیق‌ترین روش)
    2. جستجوی کلی در تمام ردیف‌های جدول
    3. جستجو در متن کامل سند
    """
    print(f"   📄 پردازش Word...")
    
    try:
        doc = Document(doc_path)
        print(f"   📊 جداول: {len(doc.tables)}")
        
        sequence_number = None
        revision = None
        date_obj = None
        date_str = None
        
        # بررسی تمام جداول
        for table_idx, table in enumerate(doc.tables):
            if sequence_number and revision and date_obj:
                break
            
            if len(table.rows) < 2:
                continue
            
            # ═══════════════════════════════════════════════════════
            # روش 1: استفاده از Header (اولویت اول - دقیق‌ترین روش)
            # ═══════════════════════════════════════════════════════
            header_row = table.rows[0]
            sequence_col_idx = None
            revision_col_idx = None
            date_col_idx = None
            
            # شناسایی ستون‌ها از روی header
            for col_idx, cell in enumerate(header_row.cells):
                header_text = cell.text.strip().lower()
                
                if 'sequence' in header_text and not sequence_col_idx:
                    sequence_col_idx = col_idx
                    print(f"   🎯 ستون Sequence: {col_idx + 1}")
                
                if 'revision' in header_text and not revision_col_idx:
                    revision_col_idx = col_idx
                    print(f"   🎯 ستون Revision: {col_idx + 1}")
                
                if 'date' in header_text and not date_col_idx:
                    date_col_idx = col_idx
            
            # استخراج مقادیر از ستون‌های شناسایی شده
            if len(table.rows) >= 2:
                data_row = table.rows[1]
                
                # Sequence Number
                if sequence_col_idx is not None and not sequence_number:
                    val = data_row.cells[sequence_col_idx].text.strip()
                    if val.isdigit() and len(val) == 4:
                        sequence_number = val
                        print(f"   ✅ Sequence (header): {val}")
                
                # Revision
                if revision_col_idx is not None and not revision:
                    val = data_row.cells[revision_col_idx].text.strip()
                    if re.match(r'^G\d{2}$', val):
                        revision = val
                        print(f"   ✅ Revision (header): {val}")
                
                # Date
                if date_col_idx is not None and not date_obj:
                    val = data_row.cells[date_col_idx].text.strip()
                    temp_date = parse_date(val)
                    if temp_date:
                        date_obj = temp_date
                        date_str = val
                        print(f"   ✅ Date (header): {val}")
            
            # ═══════════════════════════════════════════════════════
            # روش 2: جستجوی کلی در جدول (fallback)
            # ═══════════════════════════════════════════════════════
            if not sequence_number or not revision:
                for row_idx in range(1, len(table.rows)):
                    row = table.rows[row_idx]
                    values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    
                    for val in values:
                        if not sequence_number and val.isdigit() and len(val) == 4:
                            sequence_number = val
                            print(f"   ✅ Sequence (جدول): {val}")
                        
                        if not revision and re.match(r'^G\d{2}$', val):
                            revision = val
                            print(f"   ✅ Revision (جدول): {val}")
                        
                        if not date_obj:
                            temp_date = parse_date(val)
                            if temp_date:
                                date_obj = temp_date
                                date_str = val
                                print(f"   ✅ Date (جدول): {val}")
        
        # ═══════════════════════════════════════════════════════
        # روش 3: جستجو در متن کامل (آخرین راه)
        # ═══════════════════════════════════════════════════════
        if not sequence_number or not revision:
            print(f"   🔄 جستجو در متن کامل...")
            
            all_text = ""
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text += cell.text + " "
            
            for para in doc.paragraphs:
                all_text += para.text + " "
            
            patterns = [
                r'REWK\s+(\d{4})\s+(G\d{2})',
                r'(\d{4})\s+(G\d{2})',
                r'(\d{4})[^\d]+(G\d{2})',
            ]
            
            for pattern in patterns:
                match = re.search(pattern, all_text)
                if match:
                    if not sequence_number:
                        sequence_number = match.group(1)
                    if not revision:
                        revision = match.group(2)
                    print(f"   ✅ از الگو: {sequence_number}-{revision}")
                    break
            
            if not date_obj:
                date_match = re.search(r'(\d{1,2}-[A-Za-z]{3,9}-\d{4})', all_text)
                if date_match:
                    date_str = date_match.group(1)
                    date_obj = parse_date(date_str)
                    if date_obj:
                        print(f"   ✅ Date از متن: {date_str}")
        
        # نتیجه نهایی
        if sequence_number and revision:
            print(f"   ✅ نتیجه: {sequence_number}-{revision}")
            return {
                'sequence_number': sequence_number,
                'revision': revision,
                'date': date_obj,
                'date_str': date_str
            }
        else:
            print(f"   ❌ استخراج ناموفق")
            return None
            
    except Exception as e:
        print(f"   ❌ خطا: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def extract_from_pdf(pdf_path):
    """استخراج اطلاعات از فایل PDF"""
    if not PDF_SUPPORT:
        print(f"   ⚠️ PyPDF2 نصب نیست")
        return None
    
    print(f"   📄 پردازش PDF...")
    
    try:
        text = ""
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            if len(reader.pages) > 0:
                text = reader.pages[0].extract_text()
        
        if not text or len(text) < 50:
            if OCR_AVAILABLE:
                print(f"   🔍 استفاده از OCR...")
                images = convert_from_path(pdf_path, first_page=1, last_page=1, dpi=300)
                if images:
                    text = pytesseract.image_to_string(images[0], lang='eng')
        
        if not text:
            return None
        
        print(f"   📄 متن: {len(text)} کاراکتر")
        
        seq_rev_match = re.search(r'REWK\s+(\d{4})\s+(G\d{2})', text)
        if not seq_rev_match:
            seq_rev_match = re.search(r'(\d{4})\s+(G\d{2})', text)
        
        date_match = re.search(r'(\d{1,2}-[A-Za-z]{3,9}-\d{4})', text)
        
        if seq_rev_match:
            sequence_number = seq_rev_match.group(1)
            revision = seq_rev_match.group(2)
            date_str = date_match.group(1) if date_match else None
            date_obj = parse_date(date_str) if date_str else None
            
            print(f"   ✅ نتیجه: {sequence_number}-{revision}")
            
            return {
                'sequence_number': sequence_number,
                'revision': revision,
                'date': date_obj,
                'date_str': date_str
            }
        else:
            print(f"   ❌ الگو پیدا نشد")
            return None
            
    except Exception as e:
        print(f"   ❌ خطا: {str(e)}")
        return None

def create_excel_report(files_data, output_path):
    """ایجاد گزارش Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Weekly Reports"
    
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    
    headers = ['ردیف', 'نام فایل اصلی', 'نام فایل جدید', 'Sequence', 'Revision', 'تاریخ', 'وضعیت']
    ws.append(headers)
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border
    
    for idx, data in enumerate(files_data, start=1):
        row = [
            idx,
            data['old_name'],
            data['new_name'] if data['new_name'] else 'N/A',
            data['sequence_number'] if data['sequence_number'] else 'N/A',
            data['revision'] if data['revision'] else 'N/A',
            data['date'] if data['date'] else 'N/A',
            data['status']
        ]
        ws.append(row)
        
        row_num = idx + 1
        for col_idx, cell in enumerate(ws[row_num], start=1):
            cell.border = border
            cell.alignment = Alignment(horizontal='center', vertical='center')
            if col_idx == 6 and isinstance(cell.value, datetime):
                cell.number_format = 'DD/MM/YYYY'
    
    widths = [8, 40, 45, 15, 12, 15, 20]
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + idx)].width = width
    
    wb.save(output_path)
    print(f"\n📊 گزارش Excel: {output_path.name}")

def main():
    """تابع اصلی"""
    
    # ⚠️ مهم: مسیر پوشه را اینجا تنظیم کنید
    FOLDER = Path(r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\weekly")
    
    # برای تست:
    # FOLDER = Path.cwd()
    
    if not FOLDER.exists():
        print(f"❌ پوشه پیدا نشد: {FOLDER}")
        return
    
    print("\n" + "="*80)
    print(" 🔄 تغییر نام Maintenance Weekly Reports")
    print(" 📌 نسخه اصلاح شده - با استفاده از Header")
    print("="*80)
    print(f"📂 مسیر: {FOLDER}\n")
    
    # جمع‌آوری فایل‌ها
    pdf_files = list(FOLDER.glob('*.pdf'))
    word_files = list(FOLDER.glob('*.docx')) + list(FOLDER.glob('*.doc'))
    
    # فیلتر
    pdf_files = [f for f in pdf_files if not f.name.startswith('SJSC-GGNRSP-MADR-REWK-')]
    word_files = [f for f in word_files 
                  if not f.name.startswith('SJSC-GGNRSP-MADR-REWK-')
                  and not f.name.startswith('~')
                  and not f.name.startswith('.')]
    
    all_files = pdf_files + word_files
    
    print(f"📁 فایل‌های پیدا شده:")
    print(f"   PDF: {len(pdf_files)}")
    print(f"   Word: {len(word_files)}")
    print(f"   جمع: {len(all_files)}\n")
    
    if not all_files:
        print("❌ فایلی برای پردازش پیدا نشد!")
        return
    
    print("⚠️ این برنامه نام فایل‌ها را تغییر می‌دهد.")
    print("⚠️ فرمت جدید: SJSC-GGNRSP-MADR-REWK-[Sequence]-[Revision]")
    print("\nادامه می‌دهید؟ (y/n): ", end='')
    
    response = input().lower()
    if response != 'y':
        print("❌ لغو شد")
        return
    
    print("\n🔍 شروع پردازش...")
    print("-"*80)
    
    # پردازش فایل‌ها
    files_data = []
    
    for file_path in all_files:
        print(f"\n📄 {file_path.name}")
        
        if file_path.suffix.lower() == '.pdf':
            info = extract_from_pdf(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            info = extract_from_word(file_path)
        else:
            info = None
        
        if info and info['sequence_number'] and info['revision']:
            files_data.append({
                'path': file_path,
                'old_name': file_path.name,
                'sequence_number': info['sequence_number'],
                'revision': info['revision'],
                'date': info['date'],
                'date_str': info['date_str'],
                'new_name': None,
                'status': 'آماده'
            })
        else:
            files_data.append({
                'path': file_path,
                'old_name': file_path.name,
                'sequence_number': None,
                'revision': None,
                'date': None,
                'date_str': None,
                'new_name': None,
                'status': '❌ اطلاعات یافت نشد'
            })
    
    print("\n" + "="*80)
    print("🔢 تخصیص نام‌ها...")
    
    # گروه‌بندی
    groups = defaultdict(list)
    for data in files_data:
        if data['sequence_number'] and data['revision']:
            key = f"{data['sequence_number']}-{data['revision']}"
            groups[key].append(data)
    
    # تخصیص نام
    for key, group in groups.items():
        if len(group) == 1:
            data = group[0]
            ext = data['path'].suffix.lower()
            data['new_name'] = f"SJSC-GGNRSP-MADR-REWK-{data['sequence_number']}-{data['revision']}{ext}"
            print(f"   ✅ {key} → {data['new_name']}")
        else:
            print(f"   ⚠️ تکراری: {key}")
            for idx, data in enumerate(group, start=1):
                ext = data['path'].suffix.lower()
                data['new_name'] = f"SJSC-GGNRSP-MADR-REWK-{data['sequence_number']}-{data['revision']}_copy{idx}{ext}"
                print(f"      → {data['new_name']}")
    
    # تغییر نام
    print("\n🔄 تغییر نام...")
    print("-"*80)
    
    success = 0
    failed = 0
    
    for data in files_data:
        if data['new_name']:
            old_path = data['path']
            new_path = old_path.parent / data['new_name']
            
            try:
                if new_path.exists() and new_path != old_path:
                    data['status'] = '❌ نام تکراری'
                    failed += 1
                else:
                    old_path.rename(new_path)
                    data['status'] = '✅ موفق'
                    success += 1
                    print(f"✅ {old_path.name}")
                    print(f"   → {data['new_name']}")
            except Exception as e:
                data['status'] = f'❌ خطا: {str(e)}'
                failed += 1
                print(f"❌ {old_path.name}: {str(e)}")
    
    print("-"*80)
    
    # گزارش
    report_path = FOLDER / f"Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    create_excel_report(files_data, report_path)
    
    # خلاصه
    print(f"\n📊 نتیجه:")
    print(f"   ✅ موفق: {success}")
    print(f"   ❌ ناموفق: {failed}")
    
    pdf_success = sum(1 for d in files_data if '.pdf' in d['old_name'] and d['status'] == '✅ موفق')
    word_success = sum(1 for d in files_data if '.doc' in d['old_name'] and d['status'] == '✅ موفق')
    
    print(f"   📄 PDF: {pdf_success}")
    print(f"   📝 Word: {word_success}")
    print("="*80)
    print("\n✨ تمام شد!")

if __name__ == "__main__":
    main()
