#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import re
from pathlib import Path
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

try:
    from docx import Document
except:
    print("ERROR: pip install python-docx")
    exit(1)


def extract_from_docx(docx_path):
    """Extract from HEADER section of Word document"""
    try:
        doc = Document(docx_path)
        
        # Get HEADER from first section
        if not doc.sections:
            print(f"  ✗ No sections")
            return None, None, None
        
        header = doc.sections[0].header
        
        # Get tables from HEADER
        if not header.tables:
            print(f"  ✗ No tables in header")
            return None, None, None
        
        # First table in header = header info table
        header_table = header.tables[0]
        
        if len(header_table.rows) < 2:
            print(f"  ✗ Header table has < 2 rows")
            return None, None, None
        
        # Row 0 = Headers
        headers = [c.text.strip() for c in header_table.rows[0].cells]
        # Row 1 = Values
        values = [c.text.strip() for c in header_table.rows[1].cells]
        
        print(f"  Headers: {headers}")
        print(f"  Values: {values}")
        
        # Find Sequence Number column
        seq_col = None
        for i, h in enumerate(headers):
            if 'Sequence' in h and 'Number' in h:
                seq_col = i
                break
        
        # Find Revision column
        rev_col = None
        for i, h in enumerate(headers):
            if 'Revision' in h:
                rev_col = i
                break
        
        # Get Sequence Number
        seq = None
        if seq_col is not None and seq_col < len(values):
            seq = values[seq_col].strip()
            if seq:
                seq = seq.zfill(4)
                print(f"  ✓ Sequence: {seq}")
        
        # Get Revision
        rev = None
        if rev_col is not None and rev_col < len(values):
            rev = values[rev_col].strip().upper()
            print(f"  ✓ Revision: {rev}")
        
        # Get Date from body tables (approval table)
        date = None
        for table in doc.tables:
            if len(table.rows) >= 2:
                row0 = [c.text.strip() for c in table.rows[0].cells]
                # Check if this is approval table (has "Date" column)
                if 'Date' in row0:
                    date_col = row0.index('Date')
                    row1 = [c.text.strip() for c in table.rows[1].cells]
                    if date_col < len(row1):
                        date = row1[date_col].strip()
                        print(f"  ✓ Date: {date}")
                        break
        
        return seq, rev, date
        
    except Exception as e:
        print(f"  ✗ ERROR: {e}")
        import traceback
        traceback.print_exc()
        return None, None, None


def main():
    DIR = r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\weekly"
    EXCEL = r"D:\Sepher_Pasargad\works\Maintenace\Maintenance Report\All_Extracted\weekly_reports.xlsx"
    
    print("="*70)
    print("WEEKLY REPORTS - Extract from HEADER")
    print("="*70 + "\n")
    
    # Only Word files
    files = list(Path(DIR).glob("*.docx")) + list(Path(DIR).glob("*.doc"))
    
    print(f"Found {len(files)} Word files\n")
    
    results = []
    
    for i, f in enumerate(files, 1):
        if f.name.startswith('~$'):
            continue
        
        print(f"{i}. {f.name}")
        
        seq, rev, date = extract_from_docx(f)
        
        if seq and rev:
            new_name = f"SJSC-GGNRSP-EPWC-REWK-{seq}-{rev}{f.suffix}"
            results.append({
                'original_name': f.name,
                'new_name': new_name,
                'date': date or 'N/A',
                'old_path': f
            })
            print(f"  ✅ {new_name}\n")
        else:
            print(f"  ❌ FAILED (Seq={seq}, Rev={rev})\n")
    
    print("="*70)
    print(f"SUCCESS: {len(results)} files")
    print("="*70 + "\n")
    
    if not results:
        print("NO FILES!")
        return
    
    # Create Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Weekly Reports"
    
    ws.append(['Original', 'New', 'Date'])
    
    for c in ws[1]:
        c.fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        c.font = Font(bold=True, color='FFFFFF', size=12)
        c.alignment = Alignment(horizontal='center')
    
    for r in results:
        ws.append([r['original_name'], r['new_name'], r['date']])
    
    ws.column_dimensions['A'].width = 60
    ws.column_dimensions['B'].width = 50
    ws.column_dimensions['C'].width = 20
    
    wb.save(EXCEL)
    print(f"Excel: {EXCEL}\n")
    
    # Rename
    ans = input(f"Rename {len(results)} files? (yes/no): ").lower()
    
    if ans in ['yes', 'y']:
        ok = 0
        for r in results:
            new = r['old_path'].parent / r['new_name']
            if new.exists():
                print(f"EXISTS: {r['new_name']}")
                continue
            try:
                r['old_path'].rename(new)
                print(f"✅ {r['new_name']}")
                ok += 1
            except Exception as e:
                print(f"❌ {e}")
        
        print(f"\n✅ Renamed {ok}/{len(results)} files")
        print("\n🎉 DONE!")


if __name__ == "__main__":
    main()