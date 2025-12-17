#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import re
from pathlib import Path
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

try:
    import pdfplumber
    HAS_PDF = True
except:
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except:
    HAS_DOCX = False


def extract_from_pdf(pdf_path):
    if not HAS_PDF:
        return None, None, None
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = pdf.pages[0].extract_text()
            
            # Sequence Number
            seq = re.search(r'Sequence\s+Number\s+(\d+)', text, re.I)
            if not seq:
                seq = re.search(r'Number\s+(\d+)', text, re.I)
            if not seq:
                seq = re.search(r'REWK\s+(\d+)', text, re.I)
            if not seq:
                seq = re.search(r'\b(\d{4})\b', text)
            
            seq_num = seq.group(1).zfill(4) if seq else None
            
            # Revision
            rev = re.search(r'(G\d{2,3})', text, re.I)
            revision = rev.group(1).upper() if rev else None
            
            # Date
            date = re.search(r'(\d{1,2}[-/][A-Za-z]{3}[-/]\d{4})', text)
            date_str = date.group(1) if date else None
            
            return seq_num, revision, date_str
    except Exception as e:
        print(f"    PDF Error: {e}")
        return None, None, None


def extract_from_docx(docx_path):
    if not HAS_DOCX:
        return None, None, None
    
    try:
        doc = Document(docx_path)
        text = ""
        
        # Extract from ALL tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
        
        # Extract from ALL paragraphs
        for para in doc.paragraphs:
            text += para.text + " "
        
        # Sequence Number
        seq = re.search(r'Sequence\s+Number\s+(\d+)', text, re.I)
        if not seq:
            seq = re.search(r'Number\s+(\d+)', text, re.I)
        if not seq:
            seq = re.search(r'REWK\s+(\d+)', text, re.I)
        if not seq:
            seq = re.search(r'\b(\d{4})\b', text)
        
        seq_num = seq.group(1).zfill(4) if seq else None
        
        # Revision
        rev = re.search(r'(G\d{2,3})', text, re.I)
        revision = rev.group(1).upper() if rev else None
        
        # Date
        date = re.search(r'(\d{1,2}[-/][A-Za-z]{3}[-/]\d{4})', text)
        date_str = date.group(1) if date else None
        
        return seq_num, revision, date_str
    except Exception as e:
        print(f"    DOCX Error: {e}")
        return None, None, None


def process_files(directory_path):
    directory = Path(directory_path)
    
    if not directory.exists():
        print(f"ERROR: Directory not found: {directory_path}")
        return []
    
    results = []
    all_files = list(directory.glob("*.pdf")) + list(directory.glob("*.docx")) + list(directory.glob("*.doc"))
    
    print(f"\nFound {len(all_files)} files\n")
    
    for file_path in all_files:
        if file_path.name.startswith('~$'):
            continue
        
        print(f"Processing: {file_path.name}")
        
        if file_path.suffix.lower() == '.pdf':
            seq, rev, date = extract_from_pdf(file_path)
        else:
            seq, rev, date = extract_from_docx(file_path)
        
        if seq and rev:
            new_name = f"SJSC-GGNRSP-EPWC-REWK-{seq}-{rev}{file_path.suffix}"
            results.append({
                'original_name': file_path.name,
                'new_name': new_name,
                'date': date or 'N/A',
                'old_path': file_path
            })
            print(f"  -> {new_name}\n")
        else:
            print(f"  FAILED: Seq={seq}, Rev={rev}\n")
    
    return results


def create_excel(results, output_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Files"
    
    ws.append(['Original Filename', 'New Filename', 'Date'])
    
    for cell in ws[1]:
        cell.fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        cell.font = Font(bold=True, color='FFFFFF')
        cell.alignment = Alignment(horizontal='center')
    
    for item in results:
        ws.append([item['original_name'], item['new_name'], item['date']])
    
    ws.column_dimensions['A'].width = 50
    ws.column_dimensions['B'].width = 45
    ws.column_dimensions['C'].width = 20
    
    wb.save(output_path)
    print(f"\nExcel created: {output_path}")


def rename_files(results):
    print("\nRenaming...\n")
    for item in results:
        old = item['old_path']
        new = old.parent / item['new_name']
        
        if new.exists():
            print(f"EXISTS: {new.name}")
            continue
        
        try:
            old.rename(new)
            print(f"OK: {old.name} -> {new.name}")
        except Exception as e:
            print(f"FAIL: {old.name} - {e}")


def main():
    DIR = r"D:\Sepher_Pasargad\works\Production\Production Operation Report\Weekly"
    EXCEL = r"D:\Sepher_Pasargad\works\Production\Production Operation Report\Weekly\weekly_list.xlsx"
    
    print("="*70)
    print("WEEKLY REPORTS - Format: SJSC-GGNRSP-EPWC-REWK-####-G##")
    print("="*70)
    
    results = process_files(DIR)
    
    if not results:
        print("\nNo files processed!")
        return
    
    print(f"\n{len(results)} files ready")
    
    create_excel(results, EXCEL)
    
    ans = input("\nRename now? (yes/no): ").lower()
    if ans in ['yes', 'y']:
        rename_files(results)
        print("\nDONE!")
    else:
        print("\nNot renamed. Check Excel.")


if __name__ == "__main__":
    main()
